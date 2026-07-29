"""读取现有初稿结构"""
from docx import Document
import os

draft_path = r'D:\thailand study\26_7_23paper\05_output\发送\步骤15-投稿版初稿-DeepSeek事件与AI产业链非对称市场反应-中英双语引文增强版.docx'
flow_path = r'D:\thailand study\26_7_23paper\05_output\发送\我的实际流程操作和说明.docx'

print("=" * 80)
print("【初稿文档结构分析】")
print("=" * 80)

doc = Document(draft_path)
print(f"\n总段落数: {len(doc.paragraphs)}")
print(f"总表格数: {len(doc.tables)}")

print("\n--- 前100段内容预览 ---")
count = 0
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        count += 1
        style = para.style.name
        text = para.text[:150]
        print(f"[{i:3d}] {style:20s} | {text}")
        if count >= 100:
            break

print("\n--- 所有标题层级 ---")
for i, para in enumerate(doc.paragraphs):
    if 'Heading' in para.style.name or '标题' in para.style.name:
        print(f"[{i:3d}] {para.style.name:20s} | {para.text[:100]}")

print("\n" + "=" * 80)
print("【流程说明文档结构分析】")
print("=" * 80)

doc2 = Document(flow_path)
print(f"\n总段落数: {len(doc2.paragraphs)}")
print(f"总表格数: {len(doc2.tables)}")

print("\n--- 前60段内容预览 ---")
count = 0
for i, para in enumerate(doc2.paragraphs):
    if para.text.strip():
        count += 1
        style = para.style.name
        text = para.text[:150]
        print(f"[{i:3d}] {style:20s} | {text}")
        if count >= 60:
            break
