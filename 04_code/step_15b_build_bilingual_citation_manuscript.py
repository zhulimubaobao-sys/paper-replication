"""
步骤15B：在步骤15中文投稿初稿基础上，生成中英文题目、摘要和引言，
并把用户提供的19篇PDF文献纳入正文引用和参考文献表。

设计原则
--------
1. 绝不覆盖原始步骤15文档，只创建新的“中英双语引文增强版”。
2. 文献元数据以PDF正文及正式期刊页面可核验信息为准。
3. 不把日度事件研究表述为严格因果识别；月度DID未通过的事实继续保留。
4. 输出逐篇PDF的“来源—论断—引用位置”审计表，便于投稿前人工复核。

在 PyCharm 中可直接运行本文件。若项目位置未变化，无需修改任何参数。
"""

from __future__ import annotations

import csv
import hashlib
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# 一、项目路径：全部使用绝对路径，便于在 PyCharm 中一键运行
# ---------------------------------------------------------------------------
PROJECT_ROOT = Path(r"D:\thailand study\26_7_23paper")
SOURCE_DOCX = (
    PROJECT_ROOT
    / "05_output"
    / "发送"
    / "步骤15-投稿版初稿-DeepSeek事件与AI产业链非对称市场反应.docx"
)
OUTPUT_DOCX = (
    PROJECT_ROOT
    / "05_output"
    / "发送"
    / "步骤15-投稿版初稿-DeepSeek事件与AI产业链非对称市场反应-中英双语引文增强版.docx"
)
OUTPUT_DIR = PROJECT_ROOT / "05_output" / "revision_step15b"
TABLE_DIR = OUTPUT_DIR / "tables"
LOG_PATH = OUTPUT_DIR / "step15b_bilingual_citation_log.txt"
REFERENCE_AUDIT_CSV = TABLE_DIR / "table_step15b_reference_claim_audit.csv"

PDF_DIR = Path(
    r"D:\泰国留学上课资料\2.我的方向\7.24参考文献"
    r"\AI关注度冲击_40篇中英文文献包\pdfs_en"
)


# ---------------------------------------------------------------------------
# 二、19篇PDF及其核验后的用途
# verified_reference 使用正式发表信息；若PDF是早期稿，则在 metadata_note 说明。
# ---------------------------------------------------------------------------
REFERENCE_AUDIT = [
    {
        "pdf": "13_Bisetti_She_Zaldokas_2026_ESG_Shocks_in_Global_Supply_Chains.pdf",
        "citation_key": "Bisetti et al., 2026",
        "verified_reference": (
            "Bisetti, E., She, G., & Žaldokas, A. (2026). ESG shocks in global "
            "supply chains. The Review of Financial Studies, advance article, hhag001. "
            "https://doi.org/10.1093/rfs/hhag001"
        ),
        "supported_claim": "供应商环境与社会事件会诱发客户调整贸易关系并重配供应链。",
        "manuscript_location": "中文引言第4段；English Introduction第4段",
        "metadata_note": "用户PDF为较早版本；引用信息按2026年正式RFS论文核验。",
    },
    {
        "pdf": "13_姜富伟_孟令超_唐国豪_2021_媒体文本情绪与股票回报预测.pdf",
        "citation_key": "姜富伟等, 2021",
        "verified_reference": (
            "姜富伟, 孟令超, 唐国豪. (2021). 媒体文本情绪与股票回报预测. "
            "经济学（季刊）, 21(4), 1323–1344. "
            "https://doi.org/10.13821/j.cnki.ceq.2021.04.10"
        ),
        "supported_claim": "财经媒体文本情绪对中国股票回报具有样本内和样本外预测能力。",
        "manuscript_location": "中文引言第2段；English Introduction第2段",
        "metadata_note": "题名、作者、卷期、页码和DOI经期刊页面核验。",
    },
    {
        "pdf": "14_Pankratz_Schiller_2024_Climate_Change_and_Adaptation_in_Global_Supply_Chain_Networks.pdf",
        "citation_key": "Pankratz and Schiller, 2024",
        "verified_reference": (
            "Pankratz, N. M. C., & Schiller, C. M. (2024). Climate change and "
            "adaptation in global supply-chain networks. The Review of Financial "
            "Studies, 37(6), 1729–1777. https://doi.org/10.1093/rfs/hhad093"
        ),
        "supported_claim": "供应商所在地的气候冲击会传导至客户，并影响供应关系调整。",
        "manuscript_location": "中文引言第4段；English Introduction第4段",
        "metadata_note": "用户PDF含2022年FEDS稿；引用信息按2024年RFS正式版本核验。",
    },
    {
        "pdf": "15_Alfaro_Urena_Manelici_Vasquez_2022_Joining_Multinational_Supply_Chains.pdf",
        "citation_key": "Alfaro-Ureña et al., 2022",
        "verified_reference": (
            "Alfaro-Ureña, A., Manelici, I., & Vasquez, J. P. (2022). The effects "
            "of joining multinational supply chains: New evidence from firm-to-firm "
            "linkages. The Quarterly Journal of Economics, 137(3), 1495–1552. "
            "https://doi.org/10.1093/qje/qjac006"
        ),
        "supported_claim": "进入跨国公司供应链与供应商经营绩效的持续变化相关。",
        "manuscript_location": "中文引言第4段；English Introduction第4段",
        "metadata_note": "正式期刊卷期、页码和DOI由QJE页面核验。",
    },
    {
        "pdf": "17_Borusyak_Jaravel_Spiess_2024_Revisiting_Event_Study_Designs.pdf",
        "citation_key": "Borusyak et al., 2024",
        "verified_reference": (
            "Borusyak, K., Jaravel, X., & Spiess, J. (2024). Revisiting event-study "
            "designs: Robust and efficient estimation. The Review of Economic Studies, "
            "91(6), 3253–3285. https://doi.org/10.1093/restud/rdae007"
        ),
        "supported_claim": "错位处理和异质处理效应下，传统TWFE事件研究可能产生偏误。",
        "manuscript_location": "中文引言第6段；English Introduction第6段；研究设计",
        "metadata_note": "正式卷期、页码和DOI由Review of Economic Studies页面核验。",
    },
    {
        "pdf": "19_Carvalho_Nirei_Saito_Tahbaz_Salehi_2021_Supply_Chain_Disruptions.pdf",
        "citation_key": "Carvalho et al., 2021",
        "verified_reference": (
            "Carvalho, V. M., Nirei, M., Saito, Y. U., & Tahbaz-Salehi, A. (2021). "
            "Supply chain disruptions: Evidence from the Great East Japan earthquake. "
            "The Quarterly Journal of Economics, 136(2), 1255–1321. "
            "https://doi.org/10.1093/qje/qjaa044"
        ),
        "supported_claim": "局部冲击可沿供应链向上游和下游传播并被放大。",
        "manuscript_location": "中文引言第4段；English Introduction第4段",
        "metadata_note": "用户PDF为2020年稿；引用信息按2021年QJE正式版本核验。",
    },
    {
        "pdf": "20_Callaway_SantAnna_2021_Difference_in_Differences_with_Multiple_Time_Periods.pdf",
        "citation_key": "Callaway and Sant’Anna, 2021",
        "verified_reference": (
            "Callaway, B., & Sant’Anna, P. H. C. (2021). Difference-in-differences "
            "with multiple time periods. Journal of Econometrics, 225(2), 200–230. "
            "https://doi.org/10.1016/j.jeconom.2020.12.001"
        ),
        "supported_claim": "多期DID需要明确组别—时期效应、处理时点差异和聚合规则。",
        "manuscript_location": "中文引言第6段；English Introduction第6段；研究设计",
        "metadata_note": "正式卷期、页码和DOI由Journal of Econometrics页面核验。",
    },
    {
        "pdf": "22_Barber_Odean_2008_All_That_Glitters.pdf",
        "citation_key": "Barber and Odean, 2008",
        "verified_reference": (
            "Barber, B. M., & Odean, T. (2008). All that glitters: The effect of "
            "attention and news on the buying behavior of individual and institutional "
            "investors. The Review of Financial Studies, 21(2), 785–818. "
            "https://doi.org/10.1093/rfs/hhm079"
        ),
        "supported_claim": "注意力约束会影响投资者的股票选择和买入行为。",
        "manuscript_location": "中文引言第2段；English Introduction第2段；理论分析",
        "metadata_note": "PDF与RFS正式发表信息一致。",
    },
    {
        "pdf": "01_张誉夫_谢建国_2025_人工智能应用如何赋能企业供应链嵌入.pdf",
        "citation_key": "张誉夫和谢建国, 2025",
        "verified_reference": (
            "张誉夫, 谢建国. (2025). 人工智能应用如何赋能企业供应链嵌入？"
            "——基于共享商业关联的网络结构视角. 财经研究, 51(1), 63–77. "
            "https://doi.org/10.16538/j.cnki.jfe.20241022.301"
        ),
        "supported_claim": "人工智能应用有助于企业嵌入供应链网络，且存在产业链位置异质性。",
        "manuscript_location": "中文引言第1、4段；English Introduction第1、4段",
        "metadata_note": "题名、作者、卷期、页码和DOI经《财经研究》页面核验。",
    },
    {
        "pdf": "02_Li_et_al_2026_Dissecting_Corporate_Culture_Using_Generative_AI.pdf",
        "citation_key": "Li et al., 2026",
        "verified_reference": (
            "Li, K., Mai, F., Shen, R., Yang, C., & Zhang, T. (2026). Dissecting "
            "corporate culture using generative AI. The Review of Financial Studies, "
            "39(1), 253–296. https://doi.org/10.1093/rfs/hhaf081"
        ),
        "supported_claim": "生成式AI可用于提取公司文本中的经济信息，且相关差异会进入投资者反应。",
        "manuscript_location": "中文引言第3段；English Introduction第3段",
        "metadata_note": "PDF与2026年RFS正式卷期信息一致。",
    },
    {
        "pdf": "05_Bybee_Kelly_Manela_Xiu_2024_Business_News_and_Business_Cycles.pdf",
        "citation_key": "Bybee et al., 2024",
        "verified_reference": (
            "Bybee, L., Kelly, B., Manela, A., & Xiu, D. (2024). Business news and "
            "business cycles. The Journal of Finance, 79(5), 3105–3147. "
            "https://doi.org/10.1111/jofi.13377"
        ),
        "supported_claim": "新闻文本主题与市场和宏观状态的衡量及预测相关。",
        "manuscript_location": "中文引言第2段；English Introduction第2段",
        "metadata_note": "用户PDF为早期稿；引用信息按2024年Journal of Finance正式版本核验。",
    },
    {
        "pdf": "06_Charles_2025_Memory_Moves_Markets.pdf",
        "citation_key": "Charles, 2025",
        "verified_reference": (
            "Charles, C. (2025). Memory moves markets. The Review of Financial Studies, "
            "38(6), 1641–1686. https://doi.org/10.1093/rfs/hhae086"
        ),
        "supported_claim": "记忆与注意力会影响信息被市场重新激活后的交易和价格反应。",
        "manuscript_location": "中文引言第2段；English Introduction第2段；理论分析",
        "metadata_note": "卷期、页码和DOI由RFS正式页面核验。",
    },
    {
        "pdf": "07_Guo_2025_Earnings_Extrapolation_and_Predictable_Stock_Market_Returns.pdf",
        "citation_key": "Guo, 2025",
        "verified_reference": (
            "Guo, H. (2025). Earnings extrapolation and predictable stock market "
            "returns. The Review of Financial Studies, 38(6), 1730–1782. "
            "https://doi.org/10.1093/rfs/hhaf020"
        ),
        "supported_claim": "投资者外推行为与可预测的股票市场回报模式相关。",
        "manuscript_location": "中文引言第3段；English Introduction第3段",
        "metadata_note": "卷期、页码和DOI由RFS正式页面核验。",
    },
    {
        "pdf": "07_杨鹏等_2024_企业数字技术应用与创新效率提升.pdf",
        "citation_key": "杨鹏等, 2024",
        "verified_reference": (
            "杨鹏, 尹志锋, 孙宝文. (2024). 企业数字技术应用与创新效率提升. "
            "外国经济与管理, 46(11), 51–67. "
            "https://doi.org/10.16538/j.cnki.fem.20231129.401"
        ),
        "supported_claim": "数字技术应用通过知识获取和合作创新提升企业创新效率。",
        "manuscript_location": "中文引言第1段；English Introduction第1段",
        "metadata_note": "题名、作者、卷期、页码和DOI经期刊页面核验。",
    },
    {
        "pdf": "08_Andries_Bianchi_Huynh_Pouget_2025_Return_Predictability_Expectations_and_Investment.pdf",
        "citation_key": "Andries et al., 2025",
        "verified_reference": (
            "Andries, M., Bianchi, M., Huynh, K. K., & Pouget, S. (2025). Return "
            "predictability, expectations, and investment: Experimental evidence. "
            "The Review of Financial Studies, 38(6), 1687–1729. "
            "https://doi.org/10.1093/rfs/hhae088"
        ),
        "supported_claim": "信息集合会改变投资者的信念形成和投资决策。",
        "manuscript_location": "中文引言第3段；English Introduction第3段",
        "metadata_note": "用户PDF标注forthcoming；引用信息按2025年RFS正式版本核验。",
    },
    {
        "pdf": "09_Kirtac_Germano_2024_Sentiment_Trading_with_Large_Language_Models.pdf",
        "citation_key": "Kirtac and Germano, 2024",
        "verified_reference": (
            "Kirtac, K., & Germano, G. (2024). Sentiment trading with large language "
            "models. Finance Research Letters, 62, 105227. "
            "https://doi.org/10.1016/j.frl.2024.105227"
        ),
        "supported_claim": "大语言模型可从金融新闻中提取与收益预测相关的情绪信号。",
        "manuscript_location": "中文引言第3段；English Introduction第3段",
        "metadata_note": "卷号、文章号和DOI由Finance Research Letters页面核验。",
    },
    {
        "pdf": "10_Noy_Zhang_2023_Productivity_Effects_of_Generative_AI.pdf",
        "citation_key": "Noy and Zhang, 2023",
        "verified_reference": (
            "Noy, S., & Zhang, W. (2023). Experimental evidence on the productivity "
            "effects of generative artificial intelligence. Science, 381(6654), "
            "187–192. https://doi.org/10.1126/science.adh2586"
        ),
        "supported_claim": "生成式AI在受控实验中提高专业写作任务的生产率和质量。",
        "manuscript_location": "中文引言第1段；English Introduction第1段",
        "metadata_note": "用户PDF为工作论文；引用信息按2023年Science正式版本核验。",
    },
    {
        "pdf": "11_乔小勇等_2025_人工智能应用制造业出口企业高质量发展与生产网络溢出.pdf",
        "citation_key": "乔小勇等, 2025",
        "verified_reference": (
            "乔小勇, 李怡聪, 任文婷, 付大军. (2025). 人工智能应用、制造业出口企业"
            "高质量发展与生产网络溢出. 北京理工大学学报（社会科学版）, 27(6), "
            "149–165. https://doi.org/10.15918/j.jbitss1009-3370.2025.1786"
        ),
        "supported_claim": "人工智能应用可通过生产网络产生跨企业溢出效应。",
        "manuscript_location": "中文引言第1、4段；English Introduction第1、4段",
        "metadata_note": "卷期、页码和DOI经期刊页面核验。",
    },
    {
        "pdf": "12_姚加权等_2021_语调情绪及市场影响_基于金融情绪词典.pdf",
        "citation_key": "姚加权等, 2021",
        "verified_reference": (
            "姚加权, 冯绪, 王赞钧, 纪荣嵘, 张维. (2021). 语调、情绪及市场影响："
            "基于金融情绪词典. 管理科学学报, 24(5), 26–46. "
            "https://doi.org/10.19920/j.cnki.jmsc.2021.05.002"
        ),
        "supported_claim": "正式和非正式金融文本的语调情绪可预测收益、成交量和波动。",
        "manuscript_location": "中文引言第2段；English Introduction第2段",
        "metadata_note": "题名、作者、卷期和页码经《管理科学学报》页面核验。",
    },
]


# ---------------------------------------------------------------------------
# 三、双语题目、摘要、引言
# ---------------------------------------------------------------------------
CHINESE_TITLE = "DeepSeek模型发布与AI产业链的非对称市场反应"
CHINESE_SUBTITLE = "——基于A股上市公司的日度事件研究"
ENGLISH_TITLE = (
    "DeepSeek Model Releases and Asymmetric Market Reactions along the AI Value Chain"
)
ENGLISH_SUBTITLE = "Evidence from a Daily Event Study of Chinese A-Share Firms"

CHINESE_ABSTRACT = (
    "基础模型发布如何沿人工智能产业链被资本市场定价，仍缺少基于事件前企业经济功能分类"
    "和短窗口异常收益的系统证据。本文以DeepSeek-V3于2024年12月26日发布为核心事件，"
    "选取60家A股人工智能产业链上市公司，依据事件前公开的2023年年度报告，将企业按主要"
    "经济功能划分为上游、中游和下游，并采用市场模型估计日度异常收益与累计异常收益。以"
    "沪深300为主基准、[-1,+1]为预设主窗口，上游企业相对下游企业的累计异常收益高3.73个"
    "百分点（p=0.010；BH-FDR校正p=0.026）；将产业链位置编码为由下游至上游递增的有序"
    "变量后，每上移一个层级，累计异常收益提高1.92个百分点（p=0.004；BH-FDR校正"
    "p=0.006）。替换市场指数、扩大V3事件窗口以及使用DeepSeek-R1事件后，短窗口内的产业链"
    "梯度总体保持为正，但部分结果不能通过最保守的Bonferroni校正，且R1事件在较长窗口中"
    "出现衰减或反转。研究显示，基础模型发布后的短期异常收益具有产业链位置差异；该证据"
    "反映事件相关市场反应，不构成对长期价值创造或严格因果效应的单独识别。"
)
CHINESE_KEYWORDS = (
    "关键词：DeepSeek；AI产业链；投资者注意力；日度事件研究；累计异常收益；多重检验"
)

ENGLISH_ABSTRACT = (
    "How capital markets price foundation-model releases across the artificial "
    "intelligence value chain remains unclear. We examine the release of DeepSeek-V3 "
    "on 26 December 2024 using 60 Chinese A-share firms. To reduce outcome-driven "
    "classification, each firm is assigned to the upstream, midstream, or downstream "
    "segment according to its primary economic function documented in its publicly "
    "available 2023 annual report, all of which predate the event. Daily abnormal "
    "returns and cumulative abnormal returns (CARs) are estimated with a market model. "
    "Using the CSI 300 as the benchmark and the prespecified [-1,+1] window, upstream "
    "firms earn CARs that are 3.73 percentage points higher than those of downstream "
    "firms (p=0.010; Benjamini–Hochberg false-discovery-rate-adjusted p=0.026). When "
    "value-chain position is coded ordinally from downstream to upstream, a one-tier "
    "move upward is associated with a 1.92-percentage-point increase in CAR "
    "(p=0.004; adjusted p=0.006). The direction is generally preserved across "
    "alternative market indices, wider DeepSeek-V3 windows, and the DeepSeek-R1 "
    "event, although some estimates do not survive the conservative Bonferroni "
    "correction and the R1 pattern weakens or reverses over longer windows. These "
    "findings document heterogeneous short-run market reactions along the AI value "
    "chain. They should not be interpreted, on their own, as evidence of long-run "
    "value creation or a strictly identified causal effect."
)
ENGLISH_KEYWORDS = (
    "Keywords: DeepSeek; AI value chain; investor attention; daily event study; "
    "cumulative abnormal returns; multiple testing"
)

CHINESE_INTRO = [
    (
        "生成式人工智能正在从单一软件工具演变为跨行业通用技术。受控实验显示，生成式AI能够"
        "提高专业写作任务的效率与质量（Noy and Zhang, 2023）；中国企业层面的研究进一步发现，"
        "数字技术应用与创新效率提升相关（杨鹏等, 2024），人工智能应用还会影响企业供应链嵌入"
        "及生产网络溢出（张誉夫和谢建国, 2025；乔小勇等, 2025）。因此，基础模型能力的公开"
        "跃迁不仅改变技术预期，也可能同步改变资本市场对算力基础设施、软件平台和终端应用的"
        "需求判断。"
    ),
    (
        "市场如何吸收此类信息，首先受到注意力与文本信息处理约束。信息不完全会影响资产定价"
        "（Merton, 1987），显著新闻和极端交易信号能够进入投资者的有限选择集合（Barber and "
        "Odean, 2008）。在中国市场，财经媒体文本情绪和金融文本语调均具有收益或其他市场变量"
        "的预测能力（姜富伟等, 2021；姚加权等, 2021）；更广泛的商业新闻主题也能刻画并预测"
        "市场与商业周期状态（Bybee et al., 2024）。此外，信息被记忆和重新激活的方式会影响"
        "随后的市场反应（Charles, 2025）。这些研究共同表明，高传播度模型发布可能通过集中注意力"
        "和预期修正形成短期价格压力。"
    ),
    (
        "注意力并不意味着所有投资者作出相同推断。实验研究表明，信息集合会改变信念形成及信念"
        "向投资决策的传导（Andries et al., 2025），收益外推也可产生可预测的市场回报模式"
        "（Guo, 2025）。与此同时，大语言模型能够从金融新闻中提取与收益相关的情绪信号"
        "（Kirtac and Germano, 2024），生成式AI还可从分析师报告、电话会议和员工评论中提取"
        "具有资本市场含义的公司文化信息（Li et al., 2026）。这意味着DeepSeek事件既是技术"
        "供给冲击，也是一个高强度的信息与注意力冲击。"
    ),
    (
        "其次，冲击的经济含义取决于企业在供应链中的位置。既有研究表明，局部灾害会沿投入产出"
        "关系向上游和下游传播（Carvalho et al., 2021）；供应商气候暴露会影响客户经营表现和"
        "供应关系调整（Pankratz and Schiller, 2024），供应商环境与社会事件也会诱发客户重配"
        "国际采购关系（Bisetti et al., 2026）。企业进入跨国公司供应链后，其就业、生产率和客户"
        "结构会持续变化（Alfaro-Ureña et al., 2022）。结合中国证据，人工智能应用能够强化供应链"
        "嵌入和生产网络溢出（张誉夫和谢建国, 2025；乔小勇等, 2025）。然而，这些研究主要关注"
        "实体经营或网络传导，尚未直接回答同一基础模型发布为何会在AI产业链不同环节产生不同"
        "幅度的即时资本市场反应。"
    ),
    (
        "DeepSeek-V3提供了一个具有清晰公开日期且覆盖全产业链的技术信息事件。芯片、服务器、"
        "光模块、印制电路板和数据中心等上游企业提供训练与推理基础；软件平台、数据服务和系统"
        "集成企业承担技术适配与扩散；办公、金融、医疗和城市治理等下游企业将模型嵌入具体场景。"
        "由于订单映射、商业化周期和价值兑现路径不同，同一模型发布可能产生方向一致但幅度不同"
        "的市场反应。本文关注的核心问题是：在控制市场共同波动后，DeepSeek发布相关的短期异常"
        "收益是否沿产业链呈现系统性梯度。"
    ),
    (
        "研究设计上，本文采用标准日度事件研究（MacKinlay, 1997），而不把月度双重差分作为主"
        "识别框架。多期DID文献强调处理时点、异质处理效应、对照组构造和事件前趋势检验必须与"
        "目标参数一致（Callaway and Sant’Anna, 2021；Borusyak et al., 2024）。本文样本公司面对"
        "同一公开事件，缺少自然的未处理组；前期月度模型的时间安慰剂和事件前联合检验也未通过。"
        "据此，本文将日度异常收益解释为事件相关市场反应，并明确不将其升级为严格因果效应或"
        "长期基本面价值。"
    ),
    (
        "本文作出三方面贡献。第一，使用事件前公开的2023年年度报告识别60家样本公司的主要经济"
        "功能，并在结果分析前冻结分类，降低概念标签和结果导向分类风险。第二，在预设事件、市场"
        "指数和窗口下同时检验上游—下游差异与有序产业链梯度，直接呈现技术信息在价值链上的"
        "非对称定价。第三，同时披露原始p值、BH-FDR和Bonferroni校正，并以替代指数、不同窗口和"
        "DeepSeek-R1事件检验结果边界。本文由此提供一条可复核的产业链事件研究路径，而非对因果"
        "识别强度作超出数据支持的宣称。"
    ),
]

ENGLISH_INTRO = [
    (
        "Generative artificial intelligence is evolving from a stand-alone software "
        "tool into a general-purpose technology with implications across industries. "
        "Controlled evidence shows that generative AI can improve the speed and quality "
        "of professional writing tasks (Noy and Zhang, 2023). Firm-level evidence from "
        "China further links digital-technology adoption to innovation efficiency "
        "(Yang, Yin, and Sun, 2024), supply-chain embeddedness (Zhang and Xie, 2025), "
        "and production-network spillovers (Qiao et al., 2025). A public leap in "
        "foundation-model capability may therefore alter expected demand for computing "
        "infrastructure, software platforms, and downstream applications at the same time."
    ),
    (
        "The speed at which markets absorb such information is constrained by attention "
        "and textual-information processing. Incomplete information matters for asset "
        "pricing (Merton, 1987), while salient news and extreme trading signals enter "
        "investors' limited consideration sets (Barber and Odean, 2008). In China, media "
        "sentiment and financial tone contain predictive information for returns and "
        "other market outcomes (Jiang, Meng, and Tang, 2021; Yao et al., 2021). More "
        "broadly, business-news topics summarize and forecast market and business-cycle "
        "conditions (Bybee et al., 2024), and the way information is remembered and "
        "reactivated can shape later market responses (Charles, 2025). A highly visible "
        "model release can thus generate short-run price pressure through concentrated "
        "attention and expectation updating."
    ),
    (
        "Attention does not imply homogeneous inference. Experimental evidence shows "
        "that information sets change both belief formation and the mapping from beliefs "
        "to investment decisions (Andries et al., 2025), while earnings extrapolation "
        "can generate predictable return patterns (Guo, 2025). Large language models can "
        "extract sentiment signals from financial news that are related to returns "
        "(Kirtac and Germano, 2024), and generative AI can recover economically meaningful "
        "corporate-culture assessments from analyst reports, earnings calls, and employee "
        "reviews (Li et al., 2026). The DeepSeek release is therefore both a technology "
        "event and an unusually intense information-and-attention event."
    ),
    (
        "The economic meaning of the shock should also depend on a firm's position in the "
        "supply chain. Local disasters propagate both upstream and downstream through "
        "input-output linkages (Carvalho et al., 2021). Climate exposure at supplier "
        "locations affects customers and relationship termination (Pankratz and Schiller, "
        "2024), and suppliers' environmental and social incidents prompt global sourcing "
        "reallocation (Bisetti et al., 2026). Joining a multinational supply chain is "
        "followed by persistent changes in employment, productivity, and buyer structure "
        "(Alfaro-Ureña, Manelici, and Vasquez, 2022). Chinese evidence likewise connects "
        "AI adoption to supply-chain embeddedness and production-network spillovers "
        "(Zhang and Xie, 2025; Qiao et al., 2025). Yet this literature largely studies "
        "real operations or network transmission rather than heterogeneous stock-market "
        "responses across AI value-chain segments to the same foundation-model release."
    ),
    (
        "DeepSeek-V3 provides a technology-information event with a clearly observable "
        "release date and relevance across the AI value chain. Upstream firms supply "
        "chips, servers, optical modules, printed circuit boards, and data-center capacity. "
        "Midstream firms provide software platforms, data services, and systems integration. "
        "Downstream firms embed models in office, finance, health-care, and urban-service "
        "applications. Because order visibility, commercialization lags, and value-realization "
        "paths differ across these functions, a common model release may generate market "
        "reactions with the same direction but different magnitudes. We ask whether "
        "DeepSeek-related short-run abnormal returns display a systematic value-chain gradient "
        "after common market movements are removed."
    ),
    (
        "We use a conventional daily market-model event study (MacKinlay, 1997) rather than "
        "treating a monthly difference-in-differences regression as the primary identification "
        "strategy. Modern multi-period DiD methods require treatment timing, heterogeneous "
        "effects, comparison groups, pre-trends, and aggregation rules to match the estimand "
        "(Callaway and Sant’Anna, 2021; Borusyak, Jaravel, and Spiess, 2024). All firms in "
        "our sample face the same public event, leaving no natural untreated group, and our "
        "earlier monthly specifications fail time-placebo and joint pre-event tests. We "
        "therefore interpret abnormal returns as event-related market reactions, not as a "
        "strictly identified causal effect or proof of long-run fundamental value."
    ),
    (
        "This study makes three contributions. First, we classify 60 sample firms by their "
        "primary economic function using 2023 annual reports that were public before the "
        "event, and freeze the classification before analyzing the final results. Second, "
        "we test both an upstream-minus-downstream contrast and an ordered value-chain "
        "gradient under prespecified event, benchmark, and window choices. Third, we report "
        "raw p-values together with Benjamini–Hochberg false-discovery-rate and Bonferroni "
        "adjustments, and map the boundary of the findings using alternative indices, event "
        "windows, and the DeepSeek-R1 release. The result is a reproducible value-chain event "
        "study with explicit inferential limits."
    ),
]


REFERENCES = [
    (
        "Alfaro-Ureña, A., Manelici, I., & Vasquez, J. P. (2022). The effects of joining "
        "multinational supply chains: New evidence from firm-to-firm linkages. The Quarterly "
        "Journal of Economics, 137(3), 1495–1552. https://doi.org/10.1093/qje/qjac006"
    ),
    (
        "Andries, M., Bianchi, M., Huynh, K. K., & Pouget, S. (2025). Return predictability, "
        "expectations, and investment: Experimental evidence. The Review of Financial Studies, "
        "38(6), 1687–1729. https://doi.org/10.1093/rfs/hhae088"
    ),
    (
        "Barber, B. M., & Odean, T. (2008). All that glitters: The effect of attention and "
        "news on the buying behavior of individual and institutional investors. The Review of "
        "Financial Studies, 21(2), 785–818. https://doi.org/10.1093/rfs/hhm079"
    ),
    (
        "Benjamini, Y., & Hochberg, Y. (1995). Controlling the false discovery rate: "
        "A practical and powerful approach to multiple testing. Journal of the Royal "
        "Statistical Society: Series B, 57(1), 289–300."
    ),
    (
        "Bisetti, E., She, G., & Žaldokas, A. (2026). ESG shocks in global supply chains. "
        "The Review of Financial Studies, advance article, hhag001. "
        "https://doi.org/10.1093/rfs/hhag001"
    ),
    (
        "Borusyak, K., Jaravel, X., & Spiess, J. (2024). Revisiting event-study designs: "
        "Robust and efficient estimation. The Review of Economic Studies, 91(6), 3253–3285. "
        "https://doi.org/10.1093/restud/rdae007"
    ),
    (
        "Bybee, L., Kelly, B., Manela, A., & Xiu, D. (2024). Business news and business cycles. "
        "The Journal of Finance, 79(5), 3105–3147. https://doi.org/10.1111/jofi.13377"
    ),
    (
        "Callaway, B., & Sant’Anna, P. H. C. (2021). Difference-in-differences with multiple "
        "time periods. Journal of Econometrics, 225(2), 200–230. "
        "https://doi.org/10.1016/j.jeconom.2020.12.001"
    ),
    (
        "Carvalho, V. M., Nirei, M., Saito, Y. U., & Tahbaz-Salehi, A. (2021). Supply chain "
        "disruptions: Evidence from the Great East Japan earthquake. The Quarterly Journal of "
        "Economics, 136(2), 1255–1321. https://doi.org/10.1093/qje/qjaa044"
    ),
    (
        "Charles, C. (2025). Memory moves markets. The Review of Financial Studies, 38(6), "
        "1641–1686. https://doi.org/10.1093/rfs/hhae086"
    ),
    "DeepSeek-AI et al. (2024). DeepSeek-V3 Technical Report. arXiv:2412.19437.",
    (
        "DeepSeek-AI et al. (2025). DeepSeek-R1: Incentivizing reasoning capability in LLMs "
        "via reinforcement learning. arXiv:2501.12948."
    ),
    (
        "Guo, H. (2025). Earnings extrapolation and predictable stock market returns. "
        "The Review of Financial Studies, 38(6), 1730–1782. "
        "https://doi.org/10.1093/rfs/hhaf020"
    ),
    (
        "Kirtac, K., & Germano, G. (2024). Sentiment trading with large language models. "
        "Finance Research Letters, 62, 105227. "
        "https://doi.org/10.1016/j.frl.2024.105227"
    ),
    (
        "Li, K., Mai, F., Shen, R., Yang, C., & Zhang, T. (2026). Dissecting corporate "
        "culture using generative AI. The Review of Financial Studies, 39(1), 253–296. "
        "https://doi.org/10.1093/rfs/hhaf081"
    ),
    (
        "MacKinlay, A. C. (1997). Event studies in economics and finance. Journal of Economic "
        "Literature, 35(1), 13–39."
    ),
    (
        "Merton, R. C. (1987). A simple model of capital market equilibrium with incomplete "
        "information. The Journal of Finance, 42(3), 483–510."
    ),
    (
        "Noy, S., & Zhang, W. (2023). Experimental evidence on the productivity effects of "
        "generative artificial intelligence. Science, 381(6654), 187–192. "
        "https://doi.org/10.1126/science.adh2586"
    ),
    (
        "Pankratz, N. M. C., & Schiller, C. M. (2024). Climate change and adaptation in "
        "global supply-chain networks. The Review of Financial Studies, 37(6), 1729–1777. "
        "https://doi.org/10.1093/rfs/hhad093"
    ),
    (
        "姜富伟, 孟令超, 唐国豪. (2021). 媒体文本情绪与股票回报预测 [Media textual "
        "sentiment and Chinese stock return predictability]. 经济学（季刊）, 21(4), "
        "1323–1344. https://doi.org/10.13821/j.cnki.ceq.2021.04.10"
    ),
    (
        "乔小勇, 李怡聪, 任文婷, 付大军. (2025). 人工智能应用、制造业出口企业高质量发展"
        "与生产网络溢出 [Artificial intelligence applications, high-quality development "
        "of manufacturing exporters and production-network spillovers]. 北京理工大学学报"
        "（社会科学版）, 27(6), 149–165. "
        "https://doi.org/10.15918/j.jbitss1009-3370.2025.1786"
    ),
    (
        "杨鹏, 尹志锋, 孙宝文. (2024). 企业数字技术应用与创新效率提升 [Corporate "
        "digital-technology adoption and innovation-efficiency improvement]. 外国经济与管理, "
        "46(11), 51–67. https://doi.org/10.16538/j.cnki.fem.20231129.401"
    ),
    (
        "姚加权, 冯绪, 王赞钧, 纪荣嵘, 张维. (2021). 语调、情绪及市场影响：基于金融"
        "情绪词典 [Tone, sentiment, and market impact: Evidence from financial sentiment "
        "dictionaries]. 管理科学学报, 24(5), 26–46. "
        "https://doi.org/10.19920/j.cnki.jmsc.2021.05.002"
    ),
    (
        "张誉夫, 谢建国. (2025). 人工智能应用如何赋能企业供应链嵌入？——基于共享商业"
        "关联的网络结构视角 [How do AI applications empower enterprise supply-chain "
        "embedding?]. 财经研究, 51(1), 63–77. "
        "https://doi.org/10.16538/j.cnki.jfe.20241022.301"
    ),
]


def sha256_file(path: Path) -> str:
    """计算文件SHA-256，用于证明原稿未被覆盖。"""
    digest = hashlib.sha256()
    with path.open("rb") as file_handle:
        for block in iter(lambda: file_handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def set_run_font(run, name: str, size: float, bold: bool | None = None) -> None:
    """同时设置西文字体和东亚字体，减少中英文混排时的字体漂移。"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    font: str,
    size: float,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent: bool = True,
) -> None:
    """替换段落文字并应用统一格式。"""
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold)
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.25
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Pt(size * 2)
    else:
        paragraph.paragraph_format.first_line_indent = None


def insert_paragraph_before(
    target: Paragraph,
    text: str,
    *,
    style: str | None = None,
    font: str = "宋体",
    size: float = 10.5,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent: bool = True,
) -> Paragraph:
    """在目标段落之前插入新段落，并返回新段落对象。"""
    new_p = OxmlElement("w:p")
    target._p.addprevious(new_p)
    paragraph = Paragraph(new_p, target._parent)
    if style:
        paragraph.style = style
    set_paragraph_text(
        paragraph,
        text,
        font=font,
        size=size,
        bold=bold,
        alignment=alignment,
        first_line_indent=first_line_indent,
    )
    return paragraph


def delete_paragraph(paragraph: Paragraph) -> None:
    """删除指定段落节点；仅用于新建文档的内存副本，不删除任何磁盘文件。"""
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_paragraph(doc: Document, exact_text: str) -> Paragraph:
    """按完整文本查找段落，找不到即终止，避免误改位置。"""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"未找到段落：{exact_text}")


def paragraph_index(doc: Document, target: Paragraph) -> int:
    """按底层XML节点定位段落，避免python-docx包装对象比较失败。"""
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph._p is target._p:
            return index
    raise ValueError(f"段落已不在文档中：{target.text}")


def write_reference_audit() -> None:
    """输出19篇PDF的引用—论断对应表。"""
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    fieldnames = [
        "序号",
        "PDF文件名",
        "PDF完整路径",
        "正文引用键",
        "核验后参考文献",
        "该文献支持的论断",
        "稿件引用位置",
        "元数据核验说明",
        "PDF存在",
    ]
    with REFERENCE_AUDIT_CSV.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for index, item in enumerate(REFERENCE_AUDIT, start=1):
            pdf_path = PDF_DIR / item["pdf"]
            writer.writerow(
                {
                    "序号": index,
                    "PDF文件名": item["pdf"],
                    "PDF完整路径": str(pdf_path),
                    "正文引用键": item["citation_key"],
                    "核验后参考文献": item["verified_reference"],
                    "该文献支持的论断": item["supported_claim"],
                    "稿件引用位置": item["manuscript_location"],
                    "元数据核验说明": item["metadata_note"],
                    "PDF存在": "是" if pdf_path.exists() else "否",
                }
            )


def build_document() -> tuple[str, str]:
    """创建双语引文增强版DOCX，并返回原稿和新稿的SHA-256。"""
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"步骤15原稿不存在：{SOURCE_DOCX}")
    missing_pdfs = [item["pdf"] for item in REFERENCE_AUDIT if not (PDF_DIR / item["pdf"]).exists()]
    if missing_pdfs:
        raise FileNotFoundError("以下PDF不存在：\n" + "\n".join(missing_pdfs))

    original_hash_before = sha256_file(SOURCE_DOCX)
    doc = Document(SOURCE_DOCX)

    # 1. 中英文题目：保留中文主标题，增加英文题目和英文副标题。
    set_paragraph_text(
        doc.paragraphs[0],
        CHINESE_TITLE,
        font="黑体",
        size=18,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
    )
    set_paragraph_text(
        doc.paragraphs[1],
        CHINESE_SUBTITLE,
        font="宋体",
        size=13,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
    )
    set_paragraph_text(
        doc.paragraphs[2],
        "投稿版初稿（步骤15B：中英文题目、摘要与引言；参考文献增强）",
        font="宋体",
        size=10.5,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
    )
    abstract_heading = find_paragraph(doc, "摘要")
    insert_paragraph_before(
        abstract_heading,
        ENGLISH_TITLE,
        font="Times New Roman",
        size=16,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
    )
    insert_paragraph_before(
        abstract_heading,
        ENGLISH_SUBTITLE,
        font="Times New Roman",
        size=12,
        bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
    )

    # 2. 更新中文摘要，并在中文关键词之后插入英文摘要。
    abstract_index = paragraph_index(doc, abstract_heading)
    chinese_abstract_paragraph = doc.paragraphs[abstract_index + 1]
    chinese_keywords_paragraph = doc.paragraphs[abstract_index + 2]
    set_paragraph_text(
        chinese_abstract_paragraph,
        CHINESE_ABSTRACT,
        font="宋体",
        size=10.5,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    set_paragraph_text(
        chinese_keywords_paragraph,
        CHINESE_KEYWORDS,
        font="宋体",
        size=10.5,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
    )
    chinese_intro_heading = find_paragraph(doc, "一、引言")
    insert_paragraph_before(
        chinese_intro_heading,
        "Abstract",
        style="Heading 1",
        font="Times New Roman",
        size=14,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
    )
    insert_paragraph_before(
        chinese_intro_heading,
        ENGLISH_ABSTRACT,
        font="Times New Roman",
        size=10.5,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    insert_paragraph_before(
        chinese_intro_heading,
        ENGLISH_KEYWORDS,
        font="Times New Roman",
        size=10.5,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
    )

    # 3. 删除原来的5段中文引言文字，再插入新的7段中文引言和7段英文引言。
    theory_heading = find_paragraph(doc, "二、理论分析与研究假设")
    current_paragraphs = doc.paragraphs
    intro_heading_index = paragraph_index(doc, chinese_intro_heading)
    theory_heading_index = paragraph_index(doc, theory_heading)
    for paragraph in list(current_paragraphs[intro_heading_index + 1 : theory_heading_index]):
        delete_paragraph(paragraph)

    for paragraph_text in CHINESE_INTRO:
        insert_paragraph_before(
            theory_heading,
            paragraph_text,
            font="宋体",
            size=10.5,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )

    insert_paragraph_before(
        theory_heading,
        "Introduction (English)",
        style="Heading 1",
        font="Times New Roman",
        size=14,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
    )
    for paragraph_text in ENGLISH_INTRO:
        insert_paragraph_before(
            theory_heading,
            paragraph_text,
            font="Times New Roman",
            size=10.5,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )

    # 4. 在理论与研究设计段落中补充对应来源和识别边界。
    attention_paragraph = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.startswith("有限关注意味着投资者不会同时处理全部公司信息")
    )
    set_paragraph_text(
        attention_paragraph,
        (
            "有限关注意味着投资者不会同时处理全部公司信息。具有高传播度和低理解门槛的模型"
            "发布会集中吸引市场注意，并降低投资者搜索相关公司的成本。事件发生后，新增关注"
            "可能通过买入压力、预期修正和板块联动进入价格（Merton, 1987；Barber and Odean, "
            "2008；Charles, 2025；姜富伟等, 2021；姚加权等, 2021）。不过，事件研究识别的是"
            "相对于市场模型预期的短期偏离，而不是对信息完全外生性的无条件证明。"
        ),
        font="宋体",
        size=10.5,
    )

    supply_chain_paragraph = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.startswith("模型能力提升通常首先增加对训练、推理和部署基础设施")
    )
    set_paragraph_text(
        supply_chain_paragraph,
        (
            "模型能力提升通常首先增加对训练、推理和部署基础设施的预期需求。芯片、服务器、"
            "光模块、PCB和数据中心等上游投入具有较清晰的产能与订单映射，市场更容易将模型"
            "突破转化为需求预期。中游平台和软件企业承担技术适配与扩散，其受益程度取决于"
            "产品兼容性和商业化能力。下游企业虽能获得应用叙事，但业务兑现往往依赖客户付费、"
            "场景改造和持续运营，因此短窗口反应可能较弱或分散。该分层逻辑与供应链冲击可沿"
            "网络传播、并因企业经济位置而异的既有证据一致（Carvalho et al., 2021；Pankratz "
            "and Schiller, 2024；Alfaro-Ureña et al., 2022；Bisetti et al., 2026）。"
        ),
        font="宋体",
        size=10.5,
    )

    did_diagnostic_paragraph = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.startswith("作为识别诊断，本文曾使用月度面板和双向固定效应模型")
    )
    set_paragraph_text(
        did_diagnostic_paragraph,
        (
            "作为识别诊断，本文曾使用月度面板和双向固定效应模型估计事件后交互项。修正收益"
            "口径后，基准交互项为0.0528（p<0.001），但2024年1月和2024年9月伪事件显著，"
            "事件前系数联合检验也拒绝零假设。多期DID研究指出，处理时点、异质效应、对照组"
            "与目标参数必须匹配，并需审慎检验事件前趋势（Callaway and Sant’Anna, 2021；"
            "Borusyak et al., 2024）。因此，本文不以月度模型支持主结论，而将其作为事件前"
            "趋势和同期行情影响的警示。"
        ),
        font="宋体",
        size=10.5,
    )

    # 5. 重建参考文献表。只修改新文档内存中的段落，不删除任何磁盘文件。
    references_heading = find_paragraph(doc, "参考文献")
    references_heading_index = paragraph_index(doc, references_heading)
    for paragraph in list(doc.paragraphs[references_heading_index + 1 :]):
        delete_paragraph(paragraph)
    for index, reference in enumerate(REFERENCES, start=1):
        paragraph = doc.add_paragraph()
        set_paragraph_text(
            paragraph,
            f"[{index}] {reference}",
            font="Times New Roman",
            size=9.5,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=False,
        )
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(-18)

    # 6. 加入文档属性，便于识别版本。
    doc.core_properties.title = CHINESE_TITLE
    doc.core_properties.subject = "步骤15B中英双语引文增强版"
    doc.core_properties.keywords = (
        "DeepSeek; AI产业链; event study; bilingual abstract; bilingual introduction"
    )
    doc.core_properties.comments = (
        "基于步骤15原稿创建；原稿未覆盖；19篇用户提供PDF已纳入引用审计。"
    )
    doc.core_properties.modified = datetime.now()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)

    original_hash_after = sha256_file(SOURCE_DOCX)
    if original_hash_before != original_hash_after:
        raise RuntimeError("原始步骤15文档哈希发生变化，程序已停止。")
    return original_hash_before, sha256_file(OUTPUT_DOCX)


def write_log(original_hash: str, output_hash: str) -> None:
    """记录本步骤输入、输出、引用数量和不覆盖证明。"""
    lines = [
        "=" * 88,
        "步骤15B：中英文题目、摘要与引言 + 19篇PDF参考文献增强",
        "=" * 88,
        f"运行时间：{datetime.now().isoformat(timespec='seconds')}",
        f"原始文档：{SOURCE_DOCX}",
        f"输出文档：{OUTPUT_DOCX}",
        f"原始文档SHA-256：{original_hash}",
        f"输出文档SHA-256：{output_hash}",
        f"原始文档是否被覆盖：否（输入输出路径不同，且运行前后原稿哈希一致）",
        f"用户提供PDF数量：{len(REFERENCE_AUDIT)}",
        f"参考文献总数（含原稿既有方法与事件来源）：{len(REFERENCES)}",
        f"引用—论断审计表：{REFERENCE_AUDIT_CSV}",
        "",
        "重要识别边界：",
        "1. 日度事件研究用于衡量事件相关短期异常收益，不单独证明严格因果效应。",
        "2. 月度DID的时间安慰剂和事件前联合检验未通过，未作为核心识别。",
        "3. 文件名年份与正式发表年份不一致时，参考文献采用可核验的正式发表信息。",
        "=" * 88,
    ]
    LOG_PATH.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    """主程序：检查输入、写审计表、生成新DOCX、写日志并打印结果。"""
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")

    print("=" * 88)
    print("步骤15B：生成中英双语引文增强版投稿初稿")
    print("=" * 88)
    print(f"原始文档：{SOURCE_DOCX}")
    print(f"PDF目录：{PDF_DIR}")

    write_reference_audit()
    original_hash, output_hash = build_document()
    write_log(original_hash, output_hash)

    print(f"✅ 19篇PDF引用—论断审计表：{REFERENCE_AUDIT_CSV}")
    print(f"✅ 中英双语引文增强版文档：{OUTPUT_DOCX}")
    print(f"✅ 运行日志：{LOG_PATH}")
    print(f"✅ 原始文档未覆盖，SHA-256：{original_hash}")
    print(f"✅ 新文档SHA-256：{output_hash}")
    print("=" * 88)


if __name__ == "__main__":
    main()
