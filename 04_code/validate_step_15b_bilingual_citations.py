"""
步骤15B独立核验程序。

核验内容：
1. 输入、输出、日志和引用审计表是否完整；
2. 原始步骤15文档是否未被覆盖；
3. 19篇PDF是否全部存在且全部进入引用审计；
4. 中英文题目、摘要、关键词和引言是否齐全；
5. 24条唯一参考文献是否齐全；
6. 19篇新增PDF是否均在正文中被实质引用；
7. 原稿的3张表和关键实证数字是否保留；
8. 稿件是否明确保留非严格因果的识别边界。

在 PyCharm 中直接运行即可。
"""

from __future__ import annotations

import csv
import hashlib
import sys
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(r"D:\thailand study\26_7_23paper")
SOURCE_DOCX = (
    PROJECT_ROOT
    / "05_output"
    / "发送"
    / "步骤15-投稿版初稿-DeepSeek事件与AI产业链非对称市场反应.docx"
)
OUTPUT_DOCX = (
    PROJECT_ROOT
    / "05_output"
    / "发送"
    / "步骤15-投稿版初稿-DeepSeek事件与AI产业链非对称市场反应-中英双语引文增强版.docx"
)
OUTPUT_DIR = PROJECT_ROOT / "05_output" / "revision_step15b"
REFERENCE_AUDIT_CSV = (
    OUTPUT_DIR / "tables" / "table_step15b_reference_claim_audit.csv"
)
LOG_PATH = OUTPUT_DIR / "step15b_bilingual_citation_log.txt"


def sha256_file(path: Path) -> str:
    """计算文件哈希。"""
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def result_line(ok: bool, label: str) -> str:
    """统一输出核验结果。"""
    return f"{'✅' if ok else '❌'} {label}"


def main() -> int:
    """运行全部核验，科学结论边界不通过时不伪装为程序错误。"""
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")

    print("=" * 88)
    print("步骤15B：中英双语引文增强版独立核验")
    print("=" * 88)

    required_files = [SOURCE_DOCX, OUTPUT_DOCX, REFERENCE_AUDIT_CSV, LOG_PATH]
    print("\n【文件完整性】")
    file_checks = []
    for path in required_files:
        exists = path.exists() and path.stat().st_size > 0
        file_checks.append(exists)
        print(result_line(exists, str(path)))
    if not all(file_checks):
        print("\n❌ 必要文件不完整，无法继续核验。")
        return 2

    with REFERENCE_AUDIT_CSV.open("r", encoding="utf-8-sig", newline="") as handle:
        audit_rows = list(csv.DictReader(handle))

    doc = Document(OUTPUT_DOCX)
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    nonempty_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    print("\n【来源与引用核验】")
    pdf_count_ok = len(audit_rows) == 19
    pdf_exists_ok = all(row["PDF存在"] == "是" and Path(row["PDF完整路径"]).exists() for row in audit_rows)
    audit_complete_ok = all(
        row["核验后参考文献"].strip()
        and row["该文献支持的论断"].strip()
        and row["稿件引用位置"].strip()
        for row in audit_rows
    )
    print(result_line(pdf_count_ok, f"引用审计记录数：{len(audit_rows)}/19"))
    print(result_line(pdf_exists_ok, "19篇用户提供PDF全部存在"))
    print(result_line(audit_complete_ok, "每篇PDF均有核验后书目信息、支持论断和引用位置"))

    # 中文引用键与英文引用键不完全相同，因此逐项检查正文中的代表性作者标记。
    citation_tokens = [
        "Bisetti et al., 2026",
        "姜富伟等, 2021",
        "Pankratz and Schiller, 2024",
        "Alfaro-Ureña et al., 2022",
        "Borusyak et al., 2024",
        "Carvalho et al., 2021",
        "Callaway and Sant’Anna, 2021",
        "Barber and Odean, 2008",
        "张誉夫和谢建国, 2025",
        "Li et al., 2026",
        "Bybee et al., 2024",
        "Charles, 2025",
        "Guo, 2025",
        "杨鹏等, 2024",
        "Andries et al., 2025",
        "Kirtac and Germano, 2024",
        "Noy and Zhang, 2023",
        "乔小勇等, 2025",
        "姚加权等, 2021",
    ]
    missing_citations = [token for token in citation_tokens if token not in full_text]
    citation_ok = not missing_citations
    print(result_line(citation_ok, "19篇新增PDF均在正文中出现实质引用"))
    if missing_citations:
        for token in missing_citations:
            print(f"   缺失：{token}")

    print("\n【中英文结构核验】")
    structure_tokens = [
        "DeepSeek模型发布与AI产业链的非对称市场反应",
        "DeepSeek Model Releases and Asymmetric Market Reactions along the AI Value Chain",
        "摘要",
        "关键词：",
        "Abstract",
        "Keywords:",
        "一、引言",
        "Introduction (English)",
    ]
    structure_missing = [token for token in structure_tokens if token not in full_text]
    structure_ok = not structure_missing
    print(result_line(structure_ok, "中英文题目、摘要、关键词与引言结构完整"))
    if structure_missing:
        for token in structure_missing:
            print(f"   缺失：{token}")

    english_intro_word_count = 0
    in_english_intro = False
    for text in nonempty_paragraphs:
        if text == "Introduction (English)":
            in_english_intro = True
            continue
        if text == "二、理论分析与研究假设":
            in_english_intro = False
        if in_english_intro:
            english_intro_word_count += len(text.split())
    english_intro_ok = english_intro_word_count >= 650
    print(result_line(english_intro_ok, f"英文引言词数：{english_intro_word_count}（标准≥650）"))

    print("\n【参考文献与原稿保留核验】")
    references_heading_index = next(
        index for index, text in enumerate(nonempty_paragraphs) if text == "参考文献"
    )
    reference_paragraphs = nonempty_paragraphs[references_heading_index + 1 :]
    reference_count_ok = len(reference_paragraphs) == 24
    doi_count = full_text.count("https://doi.org/")
    doi_ok = doi_count >= 17
    tables_ok = len(doc.tables) == 3
    key_numbers = ["3.73个百分点", "1.92个百分点", "p=0.010", "p=0.004"]
    numbers_ok = all(token in full_text for token in key_numbers)
    boundary_tokens = [
        "不构成对长期价值创造或严格因果效应的单独识别",
        "不将其升级为严格因果效应",
        "月度DID",
    ]
    boundary_ok = all(token in full_text for token in boundary_tokens)
    print(result_line(reference_count_ok, f"唯一参考文献数量：{len(reference_paragraphs)}/24"))
    print(result_line(doi_ok, f"DOI链接数量：{doi_count}（标准≥17）"))
    print(result_line(tables_ok, f"原稿表格数量保留：{len(doc.tables)}/3"))
    print(result_line(numbers_ok, "主规格关键数字保留"))
    print(result_line(boundary_ok, "非严格因果与月度DID失败边界明确保留"))

    print("\n【不覆盖核验】")
    paths_different = SOURCE_DOCX.resolve() != OUTPUT_DOCX.resolve()
    log_text = LOG_PATH.read_text(encoding="utf-8")
    source_hash = sha256_file(SOURCE_DOCX)
    hash_recorded = source_hash in log_text
    print(result_line(paths_different, "输入与输出路径不同"))
    print(result_line(hash_recorded, "当前原稿SHA-256与构建日志记录一致"))

    all_ok = all(
        [
            *file_checks,
            pdf_count_ok,
            pdf_exists_ok,
            audit_complete_ok,
            citation_ok,
            structure_ok,
            english_intro_ok,
            reference_count_ok,
            doi_ok,
            tables_ok,
            numbers_ok,
            boundary_ok,
            paths_different,
            hash_recorded,
        ]
    )

    print("\n【最终判定】")
    if all_ok:
        print("✅ 步骤15B通过结构、引用、来源、数值保留和不覆盖核验。")
        print("⚠️ 本核验不替代人工阅读全文、目标期刊格式调整和最终语言校对。")
        return 0

    print("❌ 步骤15B核验未全部通过，请根据上方失败项修复后再投稿。")
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
