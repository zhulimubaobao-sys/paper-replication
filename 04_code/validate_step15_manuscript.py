# -*- coding: utf-8 -*-
"""步骤15投稿版Word内容一致性核验。"""

from pathlib import Path
import sys
from docx import Document

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

PATH = Path(
    r"D:/thailand study/26_7_23paper/05_output/发送/"
    r"步骤15-投稿版初稿-DeepSeek事件与AI产业链非对称市场反应.docx"
)
doc = Document(PATH)
text = "\n".join(p.text for p in doc.paragraphs)
for table in doc.tables:
    text += "\n" + "\n".join(
        "\t".join(cell.text for cell in row.cells) for row in table.rows
    )

required = {
    "新题目": "DeepSeek发布事件与AI产业链的非对称市场反应" in text,
    "60家公司": "60家" in text,
    "最终分层21/24/15":
        all(value in text for value in ("上游21家", "中游24家", "下游15家")),
    "主差异0.0373": "0.0373" in text,
    "主梯度0.0192": "0.0192" in text,
    "BH校正0.0261": "0.0261" in text,
    "Bonferroni边界0.0759": "0.0759" in text,
    "识别边界": "不能单独排除全部同期信息" in text,
    "无XX占位符": "[XX]" not in text and "[X]" not in text,
    "未宣称平行趋势成立": "平行趋势成立" not in text,
    "未宣称严格因果效应":
        "揭示了 AI 产业链上下游非对称定价的因果机制" not in text,
}

print("=" * 76)
print("步骤15投稿版内容核验")
print("=" * 76)
for name, passed in required.items():
    print(f"{'✅' if passed else '❌'} {name}")
print(f"\n段落数：{len(doc.paragraphs)}")
print(f"表格数：{len(doc.tables)}")
print("\n【最终判定】")
if all(required.values()):
    print("✅ 核心数字、分类、主张边界和占位符核验通过。")
else:
    print("❌ 内容一致性核验未通过。")
raise SystemExit(0 if all(required.values()) else 1)
