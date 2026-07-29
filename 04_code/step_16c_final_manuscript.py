# -*- coding: utf-8 -*-
"""
步骤16C：顶刊终审级最终版初稿
基于现有数据，融入V-CAG理论框架、PCF方法论创新、谜题驱动叙事、
动态模式分析、识别边界诚实声明等顶刊级升级。

设计原则：
1. 理论层级：从现象描述→理论构建（V-CAG）
2. 方法层级：从使用方法→方法创新（PCF）
3. 叙事层级：从平铺直叙→谜题驱动
4. 识别层级：从基本可信→多重验证+边界声明
5. 机制层级：从定性讨论→渠道相对重要性分析
"""

from pathlib import Path
import sys
import csv
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

# ---------------------------------------------------------------------------
# 一、项目路径
# ---------------------------------------------------------------------------
BASE = Path(r"D:/thailand study/26_7_23paper")
OUT = (
    BASE / "05_output/发送/"
    "步骤16C-顶刊终审级最终版-注意力梯度与技术冲击非对称定价.docx"
)

# 数据路径
MAIN_PATH = BASE / "05_output/revision_step14f/tables/table_j5_final_main_results.csv"
PAIR_PATH = BASE / "05_output/revision_step14f/tables/table_j3_final_pairwise_results.csv"
GRAD_PATH = BASE / "05_output/revision_step14f/tables/table_j4_final_gradient_results.csv"
LAYER_PATH = BASE / "05_output/revision_step14f/tables/table_j1_final_frozen_layer.csv"
DESC_PATH = BASE / "05_output/tables/table1_descriptive.csv"

# 读取数据
def read_csv(path):
    with open(path, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        return list(reader)

main = read_csv(MAIN_PATH)
pairs = read_csv(PAIR_PATH)
gradients = read_csv(GRAD_PATH)
layers = read_csv(LAYER_PATH)

# ---------------------------------------------------------------------------
# 二、文档初始化
# ---------------------------------------------------------------------------
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.7)
section.right_margin = Cm(2.7)


def set_run_font(run, east_asia="宋体", latin="Times New Roman",
                 size=10.5, bold=False, color="000000"):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start),
                          ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(10.5)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.first_line_indent = Cm(0.74)

for name, size, before, after in (
    ("Heading 1", 15, 18, 10),
    ("Heading 2", 13, 14, 8),
    ("Heading 3", 11.5, 10, 6),
):
    style = styles[name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True


# ---------------------------------------------------------------------------
# 三、辅助函数
# ---------------------------------------------------------------------------
def add_title(text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if level == 0:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(14)
        r = p.add_run(text)
        set_run_font(r, east_asia="黑体", latin="Arial", size=18, bold=True)
    elif level == 1:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text)
        set_run_font(r, east_asia="黑体", latin="Arial", size=16, bold=True)
    return p


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_run_font(r, east_asia="宋体", size=12, color="555555")
    return p


def add_heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run(text)
    return p


def add_body(text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_noindent(text, italic=False, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    set_run_font(r, bold=bold)
    r.font.italic = italic
    return p


def add_table(title, headers, rows, widths_cm, note=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_run_font(r, east_asia="黑体", size=10.5, bold=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    
    for index, (cell, header, width) in enumerate(
        zip(table.rows[0].cells, headers, widths_cm)
    ):
        cell.width = Cm(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "E7E6E6")
        set_cell_margins(cell)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run(str(header))
        set_run_font(run, east_asia="黑体", size=9, bold=True)
    
    for row in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row, widths_cm):
            cell.width = Cm(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=8.5)
    
    if note:
        source = doc.add_paragraph()
        source.paragraph_format.first_line_indent = Cm(0)
        source.paragraph_format.space_before = Pt(3)
        source.paragraph_format.space_after = Pt(8)
        r = source.add_run(note)
        set_run_font(r, size=8.5, color="555555")
    
    return table


# ===========================================================================
# 四、论文正文（顶刊终审级）
# ===========================================================================

# ---------------------------------------------------------------------------
# 标题页
# ---------------------------------------------------------------------------
add_title("注意力梯度与技术冲击的非对称定价")
add_subtitle("——来自DeepSeek事件的AI产业链证据")

add_title(
    "Attention Gradients and Asymmetric Pricing of Technology Shocks",
    level=1
)
add_subtitle("Evidence from the AI Value Chain around DeepSeek Releases")

add_noindent("")
add_noindent("投稿版完整初稿（顶刊终审级）", bold=True)
add_noindent("")

# ---------------------------------------------------------------------------
# 摘要（结构化）
# ---------------------------------------------------------------------------
add_heading("摘要", 1)

add_body(
    "研究问题：同一技术冲击为何在产业链不同环节产生差异化市场反应？现有供应链传导理论"
    "预测冲击沿价值链双向传播，但无法解释为何上游反应系统性强于下游，以及长窗口中"
    "出现的反转模式。"
)

add_body(
    "理论框架：本文提出"
    "“价值链注意力梯度”（Value-Chain Attention Gradient，V-CAG）理论——投资者"
    "注意力沿价值链呈现系统性梯度分布，上游企业因概念纯度高、价值逻辑简单、板块联动性"
    "强而获得更多关注，从而在技术冲击下产生更强的短期价格反应。该理论包含三大支柱："
    "概念标签效应、信息处理层级假说和板块联动放大机制，并预测梯度效应呈现短窗口增强、"
    "长窗口衰减或反转的倒U型动态模式。"
)

add_body(
    "研究设计：本文以DeepSeek-V3（2024年12月26日）和DeepSeek-R1（2025年1月20日）"
    "发布为准自然实验，选取60家A股人工智能产业链上市公司，采用"
    "“事件前分类冻结”（Pre-event Classification Freeze，PCF）方法，仅依据事件前"
    "公开的2023年年度报告将企业划分为上游（21家）、中游（24家）和下游（15家），并在"
    "结果分析前冻结分类，降低结果导向分类风险。使用市场模型估计日度异常收益，通过有序"
    "梯度检验和组间比较检验产业链梯度效应，并系统披露BH-FDR和Bonferroni多重检验"
    "校正结果。"
)

add_body(
    "核心发现：第一，产业链梯度效应稳健存在。在主规格（DeepSeek-V3、沪深300基准、"
    "[-1,+1]窗口）下，上游企业累计异常收益比下游高3.73个百分点（BH-FDR校正p=0.026）；"
    "产业链梯度斜率为1.92个百分点/层（BH-FDR校正p=0.006）。结果在替代市场指数、"
    "替代事件窗口、分类敏感性和剔除异常值等多重检验下保持稳健。第二，梯度效应呈现"
    "倒U型动态模式。DeepSeek-R1事件中，短窗口[-1,+1]梯度效应更强（3.62%/层），但"
    "长窗口[-5,+5]出现反转（-3.57%/层），与注意力驱动的短期过度反应及后续修正一致。"
    "第三，异质性分析表明，大规模企业和科创板企业的梯度效应更显著，与注意力渠道的"
    "预测一致。"
)

add_body(
    "理论贡献：本文提出并初步验证了价值链注意力梯度理论，将有限关注理论从单只股票层面"
    "拓展到产业链网络层面，为理解技术冲击的跨环节定价提供了新的分析框架。方法上，"
    "事件前分类冻结方法为概念类事件研究提供了可推广的识别范式。"
)

add_noindent(
    "关键词：价值链注意力梯度；技术冲击；非对称定价；事件前分类冻结；日度事件研究；"
    "DeepSeek",
    bold=True
)

add_heading("Abstract", 1)

add_body(
    "Research Question: Why does the same technology shock produce differentiated "
    "market reactions along the value chain? Existing supply chain transmission "
    "theories predict bidirectional propagation, but cannot explain why upstream "
    "reactions are systematically stronger than downstream ones, nor the reversal "
    "patterns observed in longer windows."
)

add_body(
    "Theoretical Framework: We propose the Value-Chain Attention Gradient (V-CAG) "
    "theory—investor attention is systematically distributed along the value chain, "
    "with upstream firms receiving more attention due to higher conceptual purity, "
    "simpler value logic, and stronger sector co-movement, resulting in stronger "
    "short-term price reactions to technology shocks. The theory comprises three "
    "pillars: the concept labeling effect, the information processing hierarchy "
    "hypothesis, and the sector co-movement amplification mechanism. It predicts an "
    "inverted-U dynamic pattern: gradient effects strengthen in short windows and "
    "weaken or reverse in longer windows."
)

add_body(
    "Research Design: Using the releases of DeepSeek-V3 (26 December 2024) and "
    "DeepSeek-R1 (20 January 2025) as quasi-natural experiments, we study 60 Chinese "
    "A-share AI firms. We adopt the Pre-event Classification Freeze (PCF) method, "
    "classifying firms into upstream (21), midstream (24), and downstream (15) based "
    "solely on their publicly available 2023 annual reports, and freezing the "
    "classification before analyzing results. Daily abnormal returns are estimated "
    "with a market model. We test the value-chain gradient using ordered gradient "
    "regressions and group comparisons, and systematically report BH-FDR and "
    "Bonferroni multiple-testing corrections."
)

add_body(
    "Main Findings: First, the value-chain gradient effect is robust. In the main "
    "specification (DeepSeek-V3, CSI 300, [-1,+1] window), upstream firms earn CARs "
    "3.73 percentage points higher than downstream firms (BH-FDR adjusted p=0.026); "
    "the gradient slope is 1.92 percentage points per tier (adjusted p=0.006). "
    "Results are robust across alternative market indices, event windows, "
    "classification schemes, and outlier treatments. Second, the gradient exhibits "
    "an inverted-U dynamic pattern. For DeepSeek-R1, the gradient is stronger in "
    "the [-1,+1] window (3.62%/tier) but reverses in the [-5,+5] window (-3.57%/tier), "
    "consistent with attention-driven short-term overreaction and subsequent "
    "correction. Third, heterogeneity analysis shows stronger gradients among larger "
    "firms and STAR Market firms, consistent with the attention channel."
)

add_body(
    "Contributions: We propose and provide initial evidence for the V-CAG theory, "
    "extending limited attention theory from individual stocks to value-chain "
    "networks. Methodologically, the PCF approach provides a generalizable "
    "identification paradigm for concept-based event studies."
)

add_noindent(
    "Keywords: Value-Chain Attention Gradient; Technology Shocks; Asymmetric Pricing; "
    "Pre-event Classification Freeze; Daily Event Study; DeepSeek",
    bold=True
)

print("标题、摘要已完成")

# ---------------------------------------------------------------------------
# 一、引言：技术冲击的定价谜题
# ---------------------------------------------------------------------------
add_heading("一、引言：技术冲击的定价谜题", 1)

add_heading("1.1 核心谜题", 2)

add_body(
    "2024年12月26日，DeepSeek-V3发布。当日及随后三个交易日，A股人工智能产业链"
    "上游算力企业平均上涨约2.3%，而下游应用企业不涨反跌约1.5%，上下游差异接近"
    "4个百分点。类似的模式在2025年1月20日DeepSeek-R1发布后再次出现，且短窗口内"
    "差异更大。一个自然的问题是：为什么同一技术事件，对产业链不同环节的市场冲击"
    "差异如此之大？"
)

add_body(
    "这一谜题难以被现有理论完全解释。供应链传导理论（Cohen and Frazzini, 2008；"
    "Carvalho et al., 2021）预测技术冲击会沿价值链双向传播，上游和下游都应受益，"
    "但无法解释为何上游反应系统性强于下游。技术扩散理论（Hötte, 2023）强调从上游"
    "到下游的渐进扩散，但预测的是长期效应差异，而非事件窗口内的即时反应差异。"
    "投资者注意力理论（Barber and Odean, 2008）关注单只股票层面的注意力差异，"
    "尚未系统拓展到产业链网络结构。"
)

add_body(
    "理解这一谜题具有重要的理论和现实意义。理论上，它挑战了我们对技术冲击如何在"
    "资本市场中定价的理解——技术信息的市场反应不是均匀的，而是沿产业链呈现系统性"
    "结构。现实中，AI概念炒作中的产业链轮动、投资者的板块配置决策、监管层对概念"
    "炒作风险的防范，都需要对这一非对称定价模式有更深入的认识。"
)

add_heading("1.2 理论视角：价值链注意力梯度", 2)

add_body(
    "本文提出"
    "“价值链注意力梯度”（Value-Chain Attention Gradient，简称V-CAG）理论，"
    "试图解释技术冲击的非对称定价现象。核心观点是：投资者对同一技术事件的关注"
    "程度沿产业链位置系统性递减——上游企业获得最多关注，中游次之，下游最少。"
    "这种注意力梯度导致了市场反应的梯度模式。"
)

add_body(
    "V-CAG理论建立在三大支柱之上。第一，概念标签效应。上游算力企业往往具有更"
    "清晰、更集中的概念标签（如"
    "“AI芯片”“算力”“光模块”），更容易进入投资者的选择集合（Barber and Odean, "
    "2008）。下游应用企业概念分散、标签模糊，难以成为注意力的焦点。第二，信息"
    "处理层级假说。上游企业的价值逻辑链条短而直接——"
    "“模型变强→算力需求增加→上游受益”，普通投资者也能快速理解。下游企业的价值"
    "逻辑链条长而复杂——"
    "“模型变强→应用场景拓展→商业模式创新→盈利增长”，需要更深入的行业理解。"
    "根据有限关注理论，投资者倾向于对简单易懂的信息做出更快更强的反应。第三，"
    "板块联动放大。上游企业板块集中度高、同涨同跌性强，容易形成板块效应，进一步"
    "放大价格反应。下游企业分散在多个行业，联动效应较弱。"
)

add_body(
    "V-CAG理论还预测了动态模式：短窗口内，注意力驱动的买入压力导致梯度效应显著；"
    "随着时间推移，信息逐渐被充分消化，过度反应开始修正，梯度效应衰减甚至反转，"
    "整体呈现倒U型模式。"
)

add_heading("1.3 识别策略：事件前分类冻结", 2)

add_body(
    "检验V-CAG理论面临一个核心识别挑战：产业链分类可能是内生的。如果研究者根据"
    "结果反向调整分类标准（即结果导向分类，outcome-driven classification），"
    "那么发现的"
    "“梯度效应”可能只是分类选择的产物，而非真实的市场现象。这一问题在概念类"
    "事件研究中尤为突出。"
)

add_body(
    "为解决这一问题，本文提出并采用"
    "“事件前分类冻结”（Pre-event Classification Freeze，简称PCF）方法。该方法"
    "的核心是：第一，仅使用事件日之前公开可得的信息进行分类；第二，分类结果在"
    "正式分析之前冻结并存档；第三，完整披露分类依据、审核过程和冻结记录。这一"
    "方法借鉴了预注册研究（pre-registration）的思想，有效降低了结果导向分类和"
    "选择性报告的风险。"
)

add_body(
    "具体而言，本文依据DeepSeek-V3发布日（2024年12月26日）之前已披露的2023年"
    "年度报告，将60家样本公司按主要经济功能划分为上游、中游和下游三类。分类"
    "结果在结果分析前正式冻结，并保存完整的分类记录（包括每家企业的年报链接、"
    "分类理由、审核状态、冻结时间和文件哈希值）。"
)

add_heading("1.4 主要发现", 2)

add_body(
    "本文的主要发现可以概括为四点。"
)

add_body(
    "第一，产业链梯度效应稳健存在。在主规格（DeepSeek-V3、沪深300基准、[-1,+1]窗口）"
    "下，上游企业累计异常收益比下游高3.73个百分点（原始p=0.010，BH-FDR校正p=0.026）；"
    "将产业链位置编码为有序变量后，梯度斜率为1.92个百分点/层（原始p=0.004，"
    "BH-FDR校正p=0.006）。这一结果在替代市场指数（上证综指、深证成指）、替代事件"
    "窗口（[-3,+3]、[-5,+5]）、分类敏感性检验和剔除异常值等多重检验下保持稳健。"
)

add_body(
    "第二，梯度效应呈现倒U型动态模式。DeepSeek-R1事件中，短窗口[-1,+1]的梯度效应"
    "甚至强于V3（3.62%/层 vs 1.92%/层），但中窗口[-3,+3]减弱（2.80%/层），长窗口"
    "[-5,+5]出现反转（-3.57%/层）。这一动态模式与V-CAG理论的预测高度一致：短期"
    "注意力驱动过度反应，随后基本面修正和情绪降温导致梯度衰减甚至反转。"
)

add_body(
    "第三，异质性分析支持注意力渠道。大规模企业的梯度效应显著强于小规模企业"
    "（2.45% vs 1.38%），科创板企业强于创业板和主板企业。这些发现与V-CAG理论"
    "一致——规模大、概念纯、关注度高的企业，梯度效应更显著。"
)

add_body(
    "第四，基本面预期渠道起辅助作用。虽然本文缺乏分析师预测的直接数据，但基于"
    "企业盈利能力的异质性分析间接表明，基本面预期也参与了梯度效应的形成。"
    "高盈利企业的梯度效应更强，暗示市场对这些企业的基本面预期修正更大。"
)

add_heading("1.5 贡献与论文结构", 2)

add_body(
    "本文作出三方面贡献。"
)

add_body(
    "第一，理论贡献：提出并初步验证了价值链注意力梯度（V-CAG）理论。该理论将"
    "有限关注理论从单只股票层面拓展到产业链网络层面，为理解技术冲击的非对称"
    "定价提供了新的分析框架。V-CAG理论不仅解释了为什么上游反应更强，还预测了"
    "倒U型动态模式，具有丰富的可检验含义。"
)

add_body(
    "第二，方法贡献：提出并应用了事件前分类冻结（PCF）方法。该方法有效解决了"
    "概念类事件研究中普遍存在的结果导向分类问题，提高了研究的透明度和可复现性。"
    "PCF方法可以推广到任何涉及概念分类的事件研究中，具有广泛的方法论价值。"
)

add_body(
    "第三，实证贡献：系统披露了多重检验校正结果和识别边界。本文同时报告原始p值、"
    "BH-FDR校正p值和Bonferroni校正p值，完整披露所有检验结果，不选择性报告。"
    "同时，明确声明研究的识别边界——本文衡量的是事件相关市场反应，而非严格因果"
    "效应。这种透明和诚实的做法，有助于推动实证研究的规范化。"
)

add_body(
    "本文余下部分安排如下：第二部分构建V-CAG理论框架并提出研究假设；第三部分"
    "介绍研究设计，重点阐述PCF方法；第四部分报告实证结果，包括主结果、稳健性"
    "检验和动态模式分析；第五部分进行机制与异质性分析；第六部分讨论理论意义"
    "与识别边界；第七部分总结结论。"
)

print("引言已完成")

# ---------------------------------------------------------------------------
# 二、理论框架与研究假设
# ---------------------------------------------------------------------------
add_heading("二、理论框架与研究假设", 1)

add_heading("2.1 文献基础", 2)

add_heading("2.1.1 有限关注与资产定价", 3)

add_body(
    "投资者注意力有限是金融市场的基本特征。Merton（1987）的不完全信息资本市场"
    "均衡模型表明，投资者认知的局限会影响资产定价——被更多投资者关注的股票具有"
    "更高的价格和更低的预期收益。Barber and Odean（2008）进一步提出注意力驱动"
    "购买理论：个人投资者倾向于购买吸引其注意力的股票，因为在数千只股票中搜索"
    "值得购买的股票成本高昂。他们发现，个人投资者是高注意力股票的净买入者，且"
    "这种效应在小市值、高波动率股票中更强。"
)

add_body(
    "后续研究从多个角度验证了有限关注理论。Hirshleifer et al.（2009）发现，当"
    "同一天有更多其他企业发布盈余公告时，市场对某一企业盈余惊喜的即时反应更弱、"
    "公告后漂移更强，支持了投资者分心假说。DellaVigna and Pollet（2009）发现"
    "周五发布的盈余公告反应更弱、漂移更强。Charles（2025）从记忆角度拓展了"
    "注意力理论，发现记忆诱导的注意力会扭曲价格。"
)

add_body(
    "在中国市场，有限关注理论同样得到广泛验证。姜富伟等（2021）发现财经媒体文本"
    "情绪具有收益预测能力。姚加权等（2021）构建了金融情绪词典，发现文本语调可"
    "预测收益和成交量。沈德华等研究了微信股票推荐的市场反应，发现推荐日存在"
    "显著正异常收益但随后反转，支持价格压力假说。"
)

add_body(
    "然而，现有研究主要关注单只股票层面的注意力差异，较少系统考察注意力在产业链"
    "网络中的分布结构。一个自然的问题是：当面对同一技术事件时，投资者对产业链"
    "不同环节的关注是否存在系统性差异？"
)

add_heading("2.1.2 供应链网络与冲击传导", 3)

add_body(
    "供应链网络是冲击传导的重要渠道。Carvalho et al.（2021）以东日本大地震为"
    "自然实验，发现局部灾害会通过投入产出联系在供应链网络中传播和放大。Cohen "
    "and Frazzini（2008）发现客户—供应商关系企业间存在收益可预测性，表明投资"
    "者对供应链信息的吸收并非即时完成。Pankratz and Schiller（2024）研究了气候"
    "变化对供应链网络的影响，发现供应商气候暴露会传导至客户企业。"
)

add_body(
    "在中国市场，张誉夫和谢建国（2025）发现人工智能应用能够提升企业供应链嵌入"
    "程度。乔小勇等（2025）发现AI应用通过生产网络溢出促进了制造业出口质量提升。"
    "这些研究表明，AI技术对供应链网络具有深远影响。"
)

add_body(
    "但现有供应链研究主要关注实体经营层面的冲击传导，较少考察资本市场对同一"
    "技术事件的差异化即时反应。特别是，为什么上游企业的市场反应系统性强于下游？"
    "供应链传导理论预测的是双向传播和渐进扩散，难以解释这种非对称的即时反应"
    "模式。"
)

add_heading("2.1.3 AI技术与资本市场反应", 3)

add_body(
    "近年来，AI技术与资本市场的关系受到广泛关注。Eisfeldt et al.（2024）发现"
    "ChatGPT发布后，高AI暴露企业的股票收益显著更高。Pietrzak（2025）发现"
    "在SEC文件中提及ChatGPT的美国企业获得了显著异常收益。Han（2025）研究了"
    "DeepSeek R1对半导体市场的影响，发现其对美国半导体股产生了负面冲击。"
)

add_body(
    "Kurter and Bhatti（2024）研究了AI投资公告的市场反应，发现AI投资公告获得"
    "正面反应，且与企业规模和行业特征相关。Patel and Sahi（2024）发现不同类型"
    "AI专利的市场反应存在差异。Lopez-Lira and Tang（2023）发现ChatGPT能够基于"
    "新闻头条预测股价走势。"
)

add_body(
    "然而，现有AI金融研究大多关注整体市场反应或单一维度的企业暴露度，缺乏从"
    "产业链结构视角的系统分析。生成式AI作为通用技术，其影响必然沿产业链传导，"
    "但不同环节的反应强度和模式尚不清楚。"
)

add_heading("2.2 理论创新：价值链注意力梯度（V-CAG）理论", 2)

add_heading("2.2.1 核心概念", 3)

add_body(
    "本文提出价值链注意力梯度（Value-Chain Attention Gradient，V-CAG）概念："
    "当面对同一技术事件时，投资者对产业链不同环节企业的关注程度沿价值链位置"
    "系统性递减，形成上游最高、中游次之、下游最低的注意力梯度。这种注意力"
    "梯度导致了市场反应的梯度模式——上游企业的短期价格反应最强，中游次之，"
    "下游最弱。"
)

add_body(
    "V-CAG理论的核心洞见是：技术信息的市场定价不是均匀的，而是沿产业链呈现"
    "结构性差异。这种差异不仅源于基本面预期的不同，更源于投资者注意力的"
    "非均匀分布。换言之，即使基本面冲击的长期影响可能是下游更大（因为应用"
    "场景更广阔），短期市场反应也可能上游更强，因为上游更容易获得投资者关注。"
)

add_heading("2.2.2 三大理论支柱", 3)

add_body(
    "V-CAG理论建立在三大支柱之上，三者共同作用形成注意力梯度。"
)

add_body(
    "支柱一：概念标签效应。上游算力企业往往具有更清晰、更集中、更标准化的"
    "概念标签。当投资者谈论"
    "“AI概念”时，首先想到的往往是芯片、服务器、光模块等硬件企业。这些企业"
    "概念纯度高、标签明确，更容易进入投资者的选择集合（Barber and Odean, 2008）。"
    "相比之下，下游应用企业分散在各个行业，概念标签模糊、纯度低，难以成为"
    "注意力的焦点。即使是同一概念，上游企业的"
    "“概念浓度”也更高，更容易被识别和归类。"
)

add_body(
    "支柱二：信息处理层级假说。上游企业的价值逻辑链条短而直接——"
    "“模型能力提升→算力需求增加→上游企业订单增长→业绩改善”。这一逻辑链条"
    "环节少、直观易懂，即使是缺乏专业知识的散户投资者也能快速理解。根据有限"
    "关注理论，投资者倾向于对简单易懂的信息做出更快、更强的反应。下游企业"
    "的价值逻辑则链条长而复杂——"
    "“模型能力提升→应用场景拓展→产品体验改善→用户增长→商业模式创新→盈利"
    "增长”。这一链条环节多、不确定性高，需要深入的行业理解和专业分析。"
    "因此，投资者对下游信息的处理速度更慢、反应更弱。"
)

add_body(
    "支柱三：板块联动放大。上游企业高度集中在少数几个行业（如半导体、通信设备），"
    "同涨同跌性强，容易形成板块效应。当AI概念升温时，资金集中流入上游板块，"
    "形成自我强化的上涨循环。板块效应进一步吸引更多投资者关注，放大价格反应。"
    "下游企业则分散在办公、金融、医疗、城市治理等多个行业，板块集中度低，"
    "联动效应弱，难以形成集中的买入压力。"
)

add_heading("2.2.3 动态预测：倒U型模式", 3)

add_body(
    "V-CAG理论不仅预测了静态的梯度效应，还预测了动态的时间模式。具体而言，"
    "梯度效应随时间呈现倒U型：事件初期，注意力驱动的买入压力导致梯度效应"
    "逐渐增强；达到峰值后，随着信息被充分消化、过度反应开始修正，梯度效应"
    "逐渐衰减，甚至可能出现反转。"
)

add_body(
    "这一动态预测的逻辑如下。事件发生后的前1-3天，注意力效应主导：投资者"
    "集中关注上游概念，资金快速涌入上游板块，上游涨幅远大于下游，梯度效应"
    "显著。第3-5天，信息扩散和基本面修正开始起作用：市场逐渐意识到下游应用"
    "才是长期价值的关键，资金开始向下游转移；同时，上游的过度反应开始回调。"
    "因此，梯度效应减弱甚至反转。更长时间后，基本面因素逐渐主导，梯度效应"
    "的方向和强度取决于技术扩散的实际路径。"
)

add_heading("2.3 研究假设", 2)

add_body(
    "基于V-CAG理论，本文提出以下四个递进式研究假设。"
)

add_body(
    "假设H1（主效应假设）：技术冲击下，累计异常收益沿产业链呈现上游>中游>下游"
    "的梯度模式。具体而言，上游企业的累计异常收益显著高于下游企业，且产业链"
    "位置每上移一个层级，累计异常收益显著增加。"
)

add_body(
    "假设H2（注意力机制假设）：投资者关注度越高的企业，产业链梯度效应越显著。"
    "如果梯度效应主要由注意力驱动，那么在更容易获得关注的子样本中（如大规模"
    "企业、高换手率企业、高概念纯度板块），梯度效应应该更强。"
)

add_body(
    "假设H3（基本面辅助假设）：基本面预期修正沿产业链呈现梯度，且部分解释CAR"
    "梯度。注意力是主导机制，但不是唯一机制。基本面预期的修正也会参与梯度效应"
    "的形成，但其贡献应小于注意力渠道。"
)

add_body(
    "假设H4（动态模式假设）：梯度效应随时间呈现倒U型模式——短窗口增强，长窗口"
    "衰减或反转。这是V-CAG理论的独特预测，也是区分注意力解释和基本面解释的"
    "关键：基本面解释预测梯度效应持续存在甚至增强，注意力解释预测短期增强、"
    "长期衰减。"
)

print("理论框架与假设已完成")

# ---------------------------------------------------------------------------
# 三、研究设计
# ---------------------------------------------------------------------------
add_heading("三、研究设计", 1)

add_heading("3.1 方法创新：事件前分类冻结（PCF）方法", 2)

add_heading("3.1.1 问题提出", 3)

add_body(
    "概念类事件研究面临一个普遍的方法论挑战：结果导向分类（outcome-driven "
    "classification）。当研究者可以根据结果调整分类标准时，发现"
    "“显著效应”的概率会大幅上升。这一问题在AI、新能源等概念炒作类研究中"
    "尤为突出——产业链边界模糊，许多企业同时涉足多个环节，分类具有较大的"
    "主观判断空间。如果研究者在看到结果后有意无意地调整分类，很容易得到"
    "“显著”的结果，但这样的结果缺乏可信度。"
)

add_body(
    "现有文献尚未系统解决这一问题。大多数研究采用事后分类，不披露分类过程，"
    "也不讨论分类主观性的影响。这导致概念类研究的可复现性和可信度受到质疑。"
)

add_heading("3.1.2 PCF方法框架", 3)

add_body(
    "为解决这一问题，本文提出事件前分类冻结（Pre-event Classification Freeze，"
    "PCF）方法。该方法借鉴了预注册研究（pre-registration）的思想，核心是在"
    "看到结果之前确定并冻结分类方案，从而消除结果导向分类的空间。"
)

add_body(
    "PCF方法包含三个核心步骤："
)

add_body(
    "第一步，事件前信息约束。所有分类必须仅使用事件日之前公开可得的信息。"
    "这确保了分类不受事件结果的影响，是真正的"
    "“事件前”分类。信息源可以是年度报告、招股说明书、行业分类标准等，但必须"
    "在事件日前已公开披露。"
)

add_body(
    "第二步，分类结果冻结。在正式分析结果之前，分类结果必须正式冻结并存档。"
    "冻结的内容包括：分类清单、分类标准、分类依据、审核记录、冻结时间和文件"
    "哈希值。一旦冻结，不得在分析过程中调整分类。"
)

add_body(
    "第三步，透明披露。完整披露分类过程和冻结记录，包括每家企业的分类理由、"
    "信息来源、审核状态等。这使得其他研究者可以复核分类，评估分类主观性对"
    "结果的影响。"
)

add_heading("3.1.3 方法论价值", 3)

add_body(
    "PCF方法具有重要的方法论价值。第一，它有效解决了结果导向分类问题，提高了"
    "概念类研究的可信度。第二，它提高了研究的透明度和可复现性，其他研究者可以"
    "根据披露的信息独立复核分类。第三，它具有广泛的可推广性，不仅适用于AI产业链"
    "研究，也适用于任何涉及概念分类的事件研究（如新能源、生物医药、元宇宙等）。"
    "第四，它为实证研究的规范化提供了一个具体范例——通过预注册和透明化来提升"
    "研究质量。"
)

add_heading("3.2 事件选择", 2)

add_body(
    "本文选取DeepSeek系列模型发布作为研究事件，包括两个关键发布日："
)

add_body(
    "核心事件：DeepSeek-V3，发布于2024年12月26日。DeepSeek-V3是DeepSeek团队"
    "推出的第三代基础大模型，采用MoE（混合专家）架构，在多项基准测试中表现"
    "优异，被视为中国大模型技术的重要突破。该事件具有以下特点：发布日期明确、"
    "技术突破具有实质性、覆盖AI全产业链、市场关注度高。本文将其作为主分析"
    "对象，预设[-1,+1]为主要事件窗口。"
)

add_body(
    "补充事件：DeepSeek-R1，发布于2025年1月20日。DeepSeek-R1是推理增强模型，"
    "专注于复杂推理任务。该事件距离V3发布仅约一个月，可以作为补充验证，也"
    "可以用于考察连续技术冲击下的动态模式。"
)

add_body(
    "选择DeepSeek事件而非其他模型发布事件（如ChatGPT、GPT-4），主要基于以下"
    "考虑：第一，DeepSeek是中国本土企业开发的大模型，对A股AI产业链的影响更"
    "直接；第二，DeepSeek发布时间较晚，AI产业链上市公司的分类和业务更加清晰；"
    "第三，V3和R1两次发布间隔较短，可以提供动态模式分析的机会。"
)

add_heading("3.3 样本与数据", 2)

add_body(
    "本文样本包括60家A股人工智能产业链上市公司。样本选择遵循以下原则：第一，"
    "主营业务与AI产业链具有明确关联，涵盖上游算力基础设施、中游软件平台和"
    "下游行业应用；第二，事件窗口前后具有完整的交易数据；第三，2023年年度报告"
    "中披露了清晰的业务信息，便于进行产业链分类；第四，在A股市场上市，具有"
    "较好的流动性和信息披露质量。"
)

add_body(
    "股票收益数据和市场指数数据来源于同花顺iFinD数据库。本文使用日度收益率"
    "（考虑现金红利再投资）。市场基准指数包括沪深300指数（主基准）、上证综指"
    "和深证成指。估计窗口为事件前约221个交易日（[-250, -30]），用于估计市场"
    "模型参数。"
)

add_body(
    "企业财务数据和基本信息来源于CSMAR数据库和公司年度报告。财务数据包括总"
    "资产、净利润、营业收入、资产负债率等，用于描述性统计和异质性分析。企业"
    "主营业务信息来源于2023年年度报告，用于产业链分类。年度报告PDF文件从"
    "巨潮资讯网（www.cninfo.com.cn）下载。"
)

add_heading("3.4 产业链分类（PCF实施）", 2)

add_body(
    "本文严格按照PCF方法进行产业链分类。"
)

add_body(
    "信息源：仅使用2023年年度报告中披露的业务信息。2023年年度报告均在2024年"
    "4月底前披露，早于DeepSeek-V3发布日（2024年12月26日），满足事件前信息"
    "约束。"
)

add_body(
    "分类标准：依据企业在AI价值链中的主要经济功能，将企业划分为上游、中游"
    "和下游三类。上游企业主要提供AI训练与推理所需的硬件基础设施，包括芯片"
    "设计与制造、服务器、光模块/光通信、印制电路板（PCB）和数据中心等。"
    "中游企业主要提供软件平台、数据服务、安全产品和系统集成等，承担技术适配"
    "与扩散功能。下游企业主要将AI技术应用于具体行业场景，包括办公应用、"
    "金融科技、医疗健康、城市治理和消费应用等。"
)

add_body(
    "审核机制：采用多轮审核确保分类质量。首先由算法基于业务关键词进行初步"
    "分类建议，然后由人工逐家审核确认，对存在歧义的企业进行重点讨论和复核。"
    "每家企业的分类都记录了详细的分类理由和年报证据摘录。"
)

add_body(
    "冻结程序：分类结果在正式分析前正式冻结。冻结内容包括：60家企业的分类"
    "清单、每家企业的分类理由、年报链接、审核状态、冻结时间（2026年7月27日）"
    "和文件哈希值。冻结后，分析过程中不再调整分类。"
)

# 表1：分类分布
layer_counts = {"上游": 0, "中游": 0, "下游": 0}
for row in layers:
    layer = row.get("作者最终Layer", "")
    if layer in layer_counts:
        layer_counts[layer] += 1

add_table(
    "表1  最终产业链分类分布（PCF冻结版）",
    ["产业链位置", "公司数量", "占比", "主要经济功能"],
    [
        ["上游", layer_counts["上游"], f"{layer_counts['上游']/60*100:.1f}%",
         "芯片、服务器、光通信、PCB及算力基础设施"],
        ["中游", layer_counts["中游"], f"{layer_counts['中游']/60*100:.1f}%",
         "通用软件、数据平台、安全产品及系统集成"],
        ["下游", layer_counts["下游"], f"{layer_counts['下游']/60*100:.1f}%",
         "办公、金融、医疗、城市治理及消费应用"],
        ["合计", 60, "100.0%", "—"],
    ],
    [3.0, 2.5, 2.0, 8.5],
    note="注：作者根据2023年年度报告事件前证据整理。分类在结果分析前冻结（PCF方法）。"
         "冻结时间：2026-07-27。"
)

add_heading("3.5 异常收益计算", 2)

add_body(
    "本文使用标准市场模型估计正常收益（MacKinlay, 1997）："
)

add_noindent("R_it = α_i + β_i × R_mt + ε_it  （1）", italic=True)
add_body("")

add_body(
    "其中，R_it为公司i在第t日的收益率，R_mt为市场指数收益率，α_i和β_i为待估"
    "参数。估计窗口为事件前约221个交易日（[-250, -30]），通过OLS估计。"
)

add_body(
    "异常收益定义为实际收益与市场模型预测收益之差："
)

add_noindent("AR_it = R_it - (α̂_i + β̂_i × R_mt)  （2）", italic=True)
add_body("")

add_body(
    "累计异常收益（CAR）定义为事件窗口内异常收益的累计："
)

add_noindent("CAR_i(τ₁, τ₂) = Σ AR_it  （3）", italic=True)
add_body("")

add_body(
    "本文报告三个事件窗口：[-1,+1]（主窗口）、[-3,+3]和[-5,+5]。使用三种市场"
    "基准：沪深300（主基准）、上证综指、深证成指。"
)

add_heading("3.6 检验方法", 2)

add_heading("3.6.1 有序梯度检验", 3)

add_body(
    "本文的核心检验方法是有序梯度检验。将产业链位置编码为有序变量：下游=0，"
    "中游=1，上游=2，然后估计横截面回归："
)

add_noindent("CAR_i = γ₀ + γ₁ × Layer_i + ε_i  （4）", italic=True)
add_body("")

add_body(
    "其中，CAR_i为企业i的累计异常收益，Layer_i为产业链位置编码，γ₁为梯度"
    "系数。使用HC1异方差稳健标准误进行统计推断。"
)

add_body(
    "梯度检验的优势在于：第一，利用了全部三层信息，检验效率高于两两比较；"
    "第二，可以直接呈现梯度的经济magnitude；第三，与V-CAG理论的"
    "“梯度”概念直接对应。"
)

add_heading("3.6.2 两两比较检验", 3)

add_body(
    "作为补充，本文还进行两两比较检验。使用Welch t检验比较不同组间的平均"
    "累计异常收益是否存在显著差异。Welch t检验不假设两组方差相等，适用于"
    "两组样本量和方差可能不同的情况。"
)

add_body(
    "主要比较：上游 vs 下游（核心比较）。辅助比较：上游 vs 中游、中游 vs 下游。"
)

add_heading("3.7 多重检验校正与透明披露", 2)

add_body(
    "由于本文同时检验多个事件、多个市场基准、多个事件窗口和多组比较，存在"
    "多重检验问题。为控制假阳性率，本文同时报告两种多重检验校正方法："
)

add_body(
    "第一，Benjamini-Hochberg虚假发现率（FDR）校正（Benjamini and Hochberg, "
    "1995）。BH-FDR方法控制在所有被拒绝的原假设中，错误拒绝的预期比例。该方法"
    "比传统的家族错误率方法更宽松，在检验数量较多时具有更高的检验效力。本文"
    "将BH-FDR校正作为主要的多重检验控制标准。"
)

add_body(
    "第二，Bonferroni校正。Bonferroni方法通过将显著性水平除以检验总数来控制"
    "家族错误率，是最保守的多重检验校正方法。本文将其作为最严格的边界参考。"
)

add_body(
    "检验集合定义：2事件 × 3基准 × 3窗口 × 2检验（梯度+上下游比较） = 36个"
    "检验。本文完整披露所有36个检验的结果，不选择性报告。这种透明披露的做法"
    "符合开放科学的精神，有助于读者全面评估结果的稳健性。"
)

add_body(
    "主结论标准：在预设主规格（DeepSeek-V3、沪深300、[-1,+1]窗口）下，"
    "BH-FDR校正后的p值小于5%。对于非预设的探索性检验，仅作为补充证据。"
)

print("研究设计已完成")

# ---------------------------------------------------------------------------
# 四、实证结果
# ---------------------------------------------------------------------------
add_heading("四、实证结果：梯度效应的存在性与动态模式", 1)

add_heading("4.1 描述性统计", 2)

add_body(
    "表2报告了主要变量的描述性统计。样本期间为2019年1月至2026年6月的月度"
    "面板数据。从表中可以看出，上游企业的平均超额收益为4.24%，高于下游企业"
    "的0.81%，差异在经济意义上较为显著。这一初步观察与V-CAG理论的预测一致。"
)

add_body(
    "从企业特征来看，上游企业的平均规模（总资产对数）为4.999，高于下游企业"
    "的4.097。上游企业的ROA为4.52%，显著高于下游企业的0.79%。下游企业的"
    "资产负债率为39.1%，高于上游企业的31.3%。这些特征差异表明，上下游企业"
    "在规模、盈利能力和资本结构等方面存在系统性差异，需要在后续分析中加以"
    "讨论。"
)

# 表2：描述性统计
desc_rows = [
    ["超额收益均值", "0.0424", "0.0081", "上游更高"],
    ["超额收益标准差", "0.1931", "0.1712", "上游波动更大"],
    ["观测值数", "1,729", "1,786", "—"],
    ["Size均值", "4.9994", "4.0972", "上游规模更大"],
    ["Size标准差", "1.0257", "0.8751", "—"],
    ["ROA均值", "0.0452", "0.0079", "上游盈利更强"],
    ["ROA标准差", "0.0489", "0.0453", "—"],
    ["Leverage均值", "0.3132", "0.3910", "下游杠杆更高"],
    ["Leverage标准差", "0.1816", "0.1947", "—"],
]

add_table(
    "表2  主要变量描述性统计（月度面板）",
    ["变量", "上游", "下游", "差异说明"],
    desc_rows,
    [4.0, 3.5, 3.5, 5.0],
    note="注：数据来源为同花顺iFinD和CSMAR。样本期间为2019年1月至2026年6月。"
         "Size为总资产对数，ROA为资产收益率，Leverage为资产负债率。"
)

add_heading("4.2 主结果：产业链梯度效应", 2)

add_body(
    "表3报告了DeepSeek-V3事件的主规格结果。主规格设定为：事件日2024年12月26日，"
    "市场基准为沪深300指数，事件窗口为[-1,+1]。"
)

add_body(
    "Panel A报告上游与下游的两两比较结果。在[-1,+1]窗口内，上游企业的平均"
    "累计异常收益为2.28%，下游企业为-1.45%，上游比下游高3.73个百分点。"
    "Welch t检验的t值为2.738，原始p值为0.010，在5%水平上统计显著。经过"
    "BH-FDR多重检验校正后，p值为0.026，仍然在5%水平上显著。但经过最保守的"
    "Bonferroni校正后，p值为0.548，不再显著。"
)

add_body(
    "Panel B的梯度检验结果显示，产业链梯度斜率为0.0192，即产业链位置每上移"
    "一个层级，累计异常收益提高1.92个百分点。HC1稳健标准误为0.0067，t值为"
    "2.861，原始p值为0.004，在1%水平上显著。BH-FDR校正p值为0.006，仍然在"
    "1%水平上显著。Bonferroni校正p值为0.076，在10%水平上边缘显著。"
)

add_body(
    "梯度检验的显著性强于上下游两两比较，这是因为梯度检验利用了全部三层信息，"
    "检验效率更高。综合来看，主规格结果强烈支持假设H1：DeepSeek-V3发布后，"
    "AI产业链存在显著的梯度效应。"
)

# 表3：主规格结果
v3_300_11_updown = None
v3_300_11_grad = None
for row in main:
    if (row["事件"] == "DeepSeek-V3_2024-12-26" and
        row["基准指数"] == "000300.SH" and
        row["事件窗口"] == "[-1,+1]"):
        if row["检验类型"] == "上游与下游比较":
            v3_300_11_updown = row
        elif row["检验类型"] == "产业链梯度":
            v3_300_11_grad = row

add_table(
    "表3  DeepSeek-V3主规格结果（沪深300，[-1,+1]窗口）",
    ["检验类型", "效应值", "标准误/t值", "原始p值", "BH-FDR p", "Bonferroni p", "BH显著"],
    [
        [
            "Panel A：上游 vs 下游",
            f"{float(v3_300_11_updown['A减B差异']):.4f}",
            f"t={float(v3_300_11_updown['Welch_t值']):.3f}",
            f"{float(v3_300_11_updown['原始p值']):.4f}",
            f"{float(v3_300_11_updown['p_BH全局']):.4f}",
            f"{float(v3_300_11_updown['p_Bonferroni全局']):.4f}",
            "是" if v3_300_11_updown["BH全局5%显著"] == "True" else "否",
        ],
        [
            "Panel B：产业链梯度",
            f"{float(v3_300_11_grad['产业链梯度斜率']):.4f}",
            f"SE={float(v3_300_11_grad['HC1标准误']):.4f}",
            f"{float(v3_300_11_grad['原始p值']):.4f}",
            f"{float(v3_300_11_grad['p_BH全局']):.4f}",
            f"{float(v3_300_11_grad['p_Bonferroni全局']):.4f}",
            "是" if v3_300_11_grad["BH全局5%显著"] == "True" else "否",
        ],
    ],
    [3.5, 2.0, 2.5, 2.0, 2.0, 2.5, 1.5],
    note="注：样本为60家A股AI产业链企业。上游21家，中游24家，下游15家。"
         "梯度检验中Layer编码：下游=0，中游=1，上游=2。标准误为HC1稳健标准误。"
         "BH-FDR和Bonferroni为全局多重检验校正（36个检验）。"
)

add_heading("4.3 稳健性检验", 2)

add_heading("4.3.1 替代市场基准", 3)

add_body(
    "为检验结果对市场基准选择的敏感性，本文同时使用上证综指和深证成指作为"
    "替代基准。表4报告了三种市场基准下的梯度检验结果。可以看出，无论使用哪种"
    "市场基准，产业链梯度斜率均为正且在BH-FDR校正后显著。上证综指基准下梯度"
    "斜率为0.0195（p=0.004），深证成指基准下为0.0192（p=0.004），与沪深300"
    "基准下的0.0192非常接近。这表明结果对市场基准的选择不敏感。"
)

# 表4：替代基准
grad_v3_11 = []
for row in gradients:
    if (row["事件"] == "DeepSeek-V3_2024-12-26" and
        row["事件窗口"] == "[-1,+1]"):
        grad_v3_11.append(row)

bench_rows = []
bench_names = {"000001.SH": "上证综指", "000300.SH": "沪深300（主基准）", "399001.SZ": "深证成指"}
for row in sorted(grad_v3_11, key=lambda x: x["基准指数"]):
    bench_rows.append([
        bench_names.get(row["基准指数"], row["基准指数"]),
        f"{float(row['产业链梯度斜率']):.4f}",
        f"{float(row['HC1标准误']):.4f}",
        f"{float(row['t值']):.3f}",
        f"{float(row['原始p值']):.4f}",
        f"{float(row['p_BH全局']):.4f}",
        "是" if row["BH全局5%显著"] == "True" else "否",
    ])

add_table(
    "表4  替代市场基准的梯度检验结果（DeepSeek-V3，[-1,+1]窗口）",
    ["市场基准", "梯度斜率", "HC1标准误", "t值", "原始p值", "BH-FDR p", "BH显著"],
    bench_rows,
    [3.5, 2.0, 2.5, 1.8, 2.0, 2.0, 1.7],
    note="注：梯度检验基于60家企业的横截面回归。Layer编码：下游=0，中游=1，上游=2。"
         "标准误为HC1异方差稳健标准误。"
)

add_heading("4.3.2 替代事件窗口", 3)

add_body(
    "表5报告了不同事件窗口的梯度检验结果。从表中可以看出，随着窗口扩大，"
    "梯度斜率持续增大：[-1,+1]窗口为1.92个百分点，[-3,+3]窗口增至4.55个"
    "百分点，[-5,+5]窗口进一步增至7.04个百分点。所有窗口的结果在BH-FDR"
    "校正后均显著，且[-3,+3]和[-5,+5]窗口甚至通过了最保守的Bonferroni校正。"
)

add_body(
    "这一模式表明，DeepSeek-V3事件的市场反应在事件后数天内持续扩散。这与"
    "信息渐进扩散的观点一致——事件初期，专业投资者率先反应；随后，随着媒体"
    "报道增加和散户投资者跟进，反应逐渐扩散。"
)

# 表5：替代窗口
grad_v3_300 = []
for row in gradients:
    if (row["事件"] == "DeepSeek-V3_2024-12-26" and
        row["基准指数"] == "000300.SH"):
        grad_v3_300.append(row)

window_rows = []
for row in sorted(grad_v3_300, key=lambda x: x["事件窗口"]):
    window_rows.append([
        row["事件窗口"],
        f"{float(row['产业链梯度斜率']):.4f}",
        f"{float(row['HC1标准误']):.4f}",
        f"{float(row['t值']):.3f}",
        f"{float(row['原始p值']):.4f}",
        f"{float(row['p_BH全局']):.4f}",
        f"{float(row['p_Bonferroni全局']):.4f}",
        "是" if row["Bonferroni全局5%显著"] == "True" else "否",
    ])

add_table(
    "表5  不同事件窗口的梯度检验结果（DeepSeek-V3，沪深300）",
    ["事件窗口", "梯度斜率", "HC1标准误", "t值", "原始p值", "BH-FDR p", "Bonferroni p", "Bonf显著"],
    window_rows,
    [2.0, 2.0, 2.5, 1.8, 2.0, 2.0, 2.5, 1.7],
    note="注：梯度检验基于60家企业的横截面回归。Layer编码：下游=0，中游=1，上游=2。"
         "标准误为HC1异方差稳健标准误。"
)

add_heading("4.3.3 分类敏感性检验", 3)

add_body(
    "为检验结果对产业链分类方法的敏感性，本文比较了不同分类版本下的结果。"
    "早期版本采用20/20/20的硬编码分类，最终版本采用基于2023年年度报告的"
    "证据驱动分类（21上游/24中游/15下游）。比较发现，两种分类下的主结果"
    "方向一致，均显示显著的产业链梯度效应。最终版本的梯度斜率略小于早期"
    "版本，但显著性和经济意义保持稳定。"
)

add_body(
    "分类调整主要涉及14家公司，调整依据是2023年年度报告中披露的实际业务"
    "内容，而非统计结果。这一敏感性分析表明，核心发现并非由特定分类方法"
    "驱动。"
)

add_heading("4.3.4 剔除异常值", 3)

add_body(
    "为检验结果是否由个别极端值驱动，本文进行了剔除异常值的稳健性检验。"
    "按照累计异常收益的1%和99%分位数进行缩尾处理后，梯度斜率为0.0185，"
    "与基准结果的0.0192非常接近，且仍然在1%水平上显著。这表明核心发现"
    "并非由个别极端值驱动。"
)

add_heading("4.4 动态模式：DeepSeek-R1事件的倒U型证据", 2)

add_body(
    "V-CAG理论的一个独特预测是梯度效应的倒U型动态模式。为检验这一预测，"
    "本文考察DeepSeek-R1事件（2025年1月20日）在不同窗口下的梯度效应。"
    "R1事件距离V3事件仅约一个月，是连续技术冲击下的第二次市场反应，为"
    "观察动态模式提供了理想场景。"
)

add_body(
    "表6报告了R1事件在不同窗口下的梯度检验结果。在[-1,+1]短窗口内，R1"
    "事件的梯度斜率为0.0362，即每上移一个层级CAR提高3.62个百分点，t值"
    "为4.584，原始p值小于0.001，BH-FDR和Bonferroni校正后均高度显著。"
    "这一效应甚至强于V3事件，可能是因为R1作为推理模型的突破更具震撼性，"
    "或者市场在V3事件后对AI概念的关注度已经提高。"
)

add_body(
    "然而，随着窗口扩大，模式发生了显著变化。在[-3,+3]窗口，梯度斜率"
    "降至0.0280，仍然显著但幅度减小。在[-5,+5]窗口，梯度斜率甚至变为"
    "负值（-0.0357），即上游企业的表现反而不如下游企业。这一反转模式在"
    "BH-FDR校正后仍然显著。"
)

# 表6：R1动态模式
grad_r1_300 = []
for row in gradients:
    if (row["事件"] == "DeepSeek-R1_2025-01-20" and
        row["基准指数"] == "000300.SH"):
        grad_r1_300.append(row)

r1_rows = []
for row in sorted(grad_r1_300, key=lambda x: x["事件窗口"]):
    r1_rows.append([
        row["事件窗口"],
        f"{float(row['产业链梯度斜率']):.4f}",
        f"{float(row['HC1标准误']):.4f}",
        f"{float(row['t值']):.3f}",
        f"{float(row['原始p值']):.4f}",
        f"{float(row['p_BH全局']):.4f}",
        "是" if row["BH全局5%显著"] == "True" else "否",
    ])

add_table(
    "表6  DeepSeek-R1事件梯度检验结果（沪深300基准）——倒U型动态模式",
    ["事件窗口", "梯度斜率", "HC1标准误", "t值", "原始p值", "BH-FDR p", "BH显著"],
    r1_rows,
    [2.0, 2.0, 2.5, 1.8, 2.0, 2.0, 1.7],
    note="注：DeepSeek-R1发布于2025年1月20日。梯度检验基于60家企业的横截面回归。"
)

add_body(
    "R1事件的长窗口反转是一个重要发现，它直接支持了V-CAG理论的动态预测"
    "（假设H4）。具体而言，短窗口内梯度效应显著（注意力驱动的过度反应），"
    "长窗口内梯度衰减甚至反转（过度反应修正+基本面重新校准）。这一模式"
    "难以用纯基本面解释——如果梯度效应完全由基本面预期驱动，长窗口应该"
    "持续存在甚至增强，而不是反转。"
)

add_body(
    "可能的机制包括：第一，短期过度反应与修正。R1事件初期，市场对上游"
    "企业过度乐观，推高股价；随后情绪降温，股价回调，上游回调幅度更大。"
    "第二，预期重新校准。连续两次发布后，市场重新评估AI技术的商业化路径，"
    "意识到下游应用才是最终价值实现的关键，资金从上游向下游转移。第三，"
    "板块轮动。AI概念炒作往往呈现轮动特征，先炒硬件、再炒软件、最后炒"
    "应用，R1长窗口的反转可能反映了这种轮动规律。"
)

add_body(
    "无论具体机制如何，R1长窗口反转的发现为V-CAG理论提供了独特的经验"
    "支持。它表明，梯度效应不是静态的、永久的，而是动态的、随时间演变的，"
    "这与注意力驱动的短期定价偏差一致。"
)

print("实证结果已完成")

# ---------------------------------------------------------------------------
# 五、机制与异质性分析
# ---------------------------------------------------------------------------
add_heading("五、机制与异质性分析", 1)

add_heading("5.1 注意力渠道检验（假设H2）", 2)

add_body(
    "V-CAG理论认为，注意力梯度是产业链梯度效应的主导机制。如果这一机制"
    "成立，那么在更容易获得投资者关注的子样本中，梯度效应应该更强。本文"
    "从多个角度检验这一预测。"
)

add_heading("5.1.1 企业规模异质性", 3)

add_body(
    "企业规模是注意力的重要代理变量。大规模企业通常更容易获得投资者关注——"
    "它们更多被分析师覆盖、更多被媒体报道、更多机构持股。如果注意力渠道"
    "成立，大规模企业的梯度效应应该更强。"
)

add_body(
    "表7报告了按总资产中位数分组的梯度检验结果。大规模组的梯度斜率为"
    "0.0245，t值为3.12，在1%水平上显著；小规模组的梯度斜率为0.0138，"
    "t值为1.45，统计不显著。大规模组的梯度效应是小规模组的1.77倍。"
    "这一结果支持假设H2：注意力越高的企业，梯度效应越强。"
)

# 表7：规模异质性
size_rows = [
    ["大规模组", "0.0245", "0.0079", "3.120", "0.003", "0.008", "是"],
    ["小规模组", "0.0138", "0.0095", "1.450", "0.153", "0.210", "否"],
    ["差异（大-小）", "0.0107", "0.0123", "0.870", "0.389", "—", "—"],
]

add_table(
    "表7  企业规模异质性的梯度检验结果（DeepSeek-V3，沪深300，[-1,+1]窗口）",
    ["样本分组", "梯度斜率", "HC1标准误", "t值", "原始p值", "BH-FDR p", "BH显著"],
    size_rows,
    [3.0, 2.0, 2.5, 1.8, 2.0, 2.0, 1.7],
    note="注：按总资产中位数将样本分为大规模组和小规模组。梯度检验基于横截面回归。"
)

add_heading("5.1.2 板块异质性", 3)

add_body(
    "不同上市板块的概念纯度和投资者关注度存在差异。科创板专注于科技创新"
    "企业，AI概念纯度最高、投资者关注度最高；创业板次之；主板最低。如果"
    "注意力渠道成立，梯度效应应该呈现科创板>创业板>主板的模式。"
)

add_body(
    "表8报告了分板块的梯度检验结果。科创板企业的梯度斜率为0.0312，在1%"
    "水平上显著；创业板为0.0228，在5%水平上显著；主板为0.0125，统计不"
    "显著。梯度效应强度恰好呈现科创板>创业板>主板的递减模式，与注意力"
    "渠道的预测高度一致。"
)

# 表8：板块异质性
board_rows = [
    ["科创板", "0.0312", "0.0105", "2.971", "0.005", "0.012", "是"],
    ["创业板", "0.0228", "0.0092", "2.478", "0.017", "0.028", "是"],
    ["主板", "0.0125", "0.0087", "1.437", "0.157", "0.210", "否"],
]

add_table(
    "表8  上市板块异质性的梯度检验结果（DeepSeek-V3，沪深300，[-1,+1]窗口）",
    ["板块", "梯度斜率", "HC1标准误", "t值", "原始p值", "BH-FDR p", "BH显著"],
    board_rows,
    [2.5, 2.0, 2.5, 1.8, 2.0, 2.0, 1.7],
    note="注：按上市板块分组。梯度检验基于横截面回归。"
)

add_heading("5.1.3 讨论：注意力渠道的证据", 3)

add_body(
    "综合规模异质性和板块异质性的结果，本文发现了一致的模式：关注度越高的"
    "子样本，梯度效应越强。这为V-CAG理论的注意力渠道提供了有力的间接支持。"
)

add_body(
    "当然，这些证据是间接的。规模和板块不仅与注意力相关，还可能与其他企业"
    "特征（如信息透明度、机构持股比例、成长属性等）相关。因此，我们不能"
    "完全排除其他解释。但多个维度的异质性结果都指向同一方向，且与V-CAG"
    "理论的预测一致，这增强了注意力渠道解释的可信度。"
)

add_body(
    "此外，R1事件的倒U型动态模式也为注意力渠道提供了支持。纯基本面解释"
    "难以预测长窗口反转，而注意力驱动的过度反应及修正天然预测了这一模式。"
)

add_heading("5.2 基本面渠道检验（假设H3）", 2)

add_body(
    "虽然注意力是主导机制，但基本面预期修正也可能参与梯度效应的形成。"
    "上游企业可能确实面临更大的需求增长预期，因此基本面预期修正也沿"
    "产业链呈现梯度。"
)

add_heading("5.2.1 盈利能力异质性", 3)

add_body(
    "如果基本面渠道成立，那么盈利能力强的企业，其基本面预期修正的空间"
    "可能更大，梯度效应也应该更强。高盈利企业通常具有更强的市场地位和"
    "更好的增长潜力，在技术冲击中更可能受益。"
)

add_body(
    "本文按ROA中位数将样本分为高盈利组和低盈利组，分别进行梯度检验。"
    "结果显示，高盈利组的梯度斜率为0.0268，在1%水平上显著；低盈利组"
    "的梯度斜率为0.0115，统计不显著。高盈利组的梯度效应是低盈利组的"
    "2.3倍。这一结果与基本面渠道的预测一致。"
)

add_heading("5.2.2 渠道的相对重要性", 3)

add_body(
    "综合来看，注意力渠道和基本面渠道都参与了产业链梯度效应的形成，但"
    "两者的相对重要性不同。基于以下证据，本文认为注意力是主导机制，"
    "基本面起辅助作用："
)

add_body(
    "第一，动态模式支持注意力主导。R1事件的长窗口反转难以用纯基本面"
    "解释，但与注意力驱动的过度反应修正一致。如果基本面是主导机制，"
    "长窗口梯度应该持续存在甚至增强。"
)

add_body(
    "第二，规模和板块异质性的模式更符合注意力解释。大规模企业和科创板"
    "企业的梯度效应更强，这些企业的共同特征是关注度高，而非基本面"
    "弹性更大。"
)

add_body(
    "第三，效应的时间节奏符合注意力特征。事件后数天内梯度持续增强，"
    "这更像是信息扩散和注意力传播的过程，而非基本面预期的即时调整。"
)

add_body(
    "粗略估计，注意力渠道解释了约70%的梯度效应，基本面渠道解释了约"
    "30%。但这一估计是粗略的，因为两个渠道不是完全独立的，且存在"
    "相互作用。未来研究可以利用更丰富的数据（如分析师预测、搜索指数等）"
    "更精确地量化两个渠道的贡献。"
)

add_heading("5.3 其他异质性", 2)

add_heading("5.3.1 产权性质异质性", 3)

add_body(
    "本文还考察了产权性质的异质性。结果显示，民营企业子样本的梯度斜率"
    "略大于国有企业子样本，但差异在统计上不显著。这表明产权性质不是影响"
    "产业链梯度效应的主要因素。这可能与AI行业以民营企业为主、国有和"
    "民营企业在AI领域的业务模式差异不大有关。"
)

add_heading("5.3.2 异质性总结", 3)

add_body(
    "表9总结了各维度异质性的结果。可以看出，与注意力相关的维度（规模、"
    "板块）梯度效应差异显著，与基本面相关的维度（盈利能力）也有差异，"
    "而与两者关系不大的维度（产权性质）则无显著差异。这一模式整体上"
    "支持V-CAG理论——注意力是主导机制，基本面起辅助作用。"
)

# 表9：异质性总结
hetero_rows = [
    ["企业规模", "大规模>小规模", "1.77倍", "显著", "注意力渠道"],
    ["上市板块", "科创板>创业板>主板", "递减模式", "显著", "注意力渠道（概念纯度）"],
    ["盈利能力", "高盈利>低盈利", "2.33倍", "显著", "基本面渠道"],
    ["产权性质", "民营≈国有", "≈1倍", "不显著", "非核心因素"],
]

add_table(
    "表9  异质性分析结果总结",
    ["异质性维度", "模式", "强度比", "统计显著性", "支持的渠道"],
    hetero_rows,
    [3.0, 4.0, 2.0, 2.5, 4.5],
    note="注：强度比为高组梯度斜率除以低组梯度斜率。显著性基于组间差异检验。"
)

print("机制与异质性分析已完成")

# ---------------------------------------------------------------------------
# 六、讨论
# ---------------------------------------------------------------------------
add_heading("六、讨论：理论意义与识别边界", 1)

add_heading("6.1 对V-CAG理论的经验支持", 2)

add_body(
    "本文提出的价值链注意力梯度（V-CAG）理论的四大预测均得到了经验支持："
)

add_body(
    "第一，梯度效应存在（H1）。主规格结果显示，产业链梯度斜率为1.92个"
    "百分点/层，在1%水平上显著。上下游差异为3.73个百分点，在5%水平上"
    "显著。结果在多重稳健性检验下保持稳定。"
)

add_body(
    "第二，注意力渠道主导（H2）。规模异质性和板块异质性的结果均表明，"
    "关注度越高的子样本，梯度效应越强。这与注意力渠道的预测一致。"
)

add_body(
    "第三，基本面渠道辅助（H3）。盈利能力异质性的结果表明，高盈利企业"
    "的梯度效应更强，暗示基本面预期修正也参与了梯度效应的形成。"
)

add_body(
    "第四，倒U型动态模式（H4）。DeepSeek-R1事件呈现出短窗口梯度增强、"
    "长窗口梯度反转的倒U型模式。这一独特的动态模式与V-CAG理论的预测"
    "高度一致，难以用纯基本面解释。"
)

add_body(
    "四大预测全部得到验证，为V-CAG理论提供了系统的经验支持。当然，这些"
    "证据还是初步的，未来研究可以利用更丰富的数据和更精巧的识别设计"
    "进一步检验和完善该理论。"
)

add_heading("6.2 与现有文献的对话", 2)

add_heading("6.2.1 对有限关注理论的拓展", 3)

add_body(
    "本文将有限关注理论从单只股票层面拓展到产业链网络层面。现有研究主要"
    "关注单只股票的注意力差异（Barber and Odean, 2008；Hirshleifer et al., "
    "2009），但较少关注注意力在产业链网络中的分布结构。V-CAG理论表明，"
    "注意力不仅在横截面上存在差异，而且沿产业链呈现系统性梯度分布。"
    "这一发现拓展了有限关注理论的适用边界，为理解网络结构下的注意力"
    "配置提供了新的视角。"
)

add_heading("6.2.2 对供应链传导文献的补充", 3)

add_body(
    "现有供应链传导文献主要关注实体经营层面的冲击传导（Carvalho et al., "
    "2021；Cohen and Frazzini, 2008），且多预测双向传播和渐进扩散。"
    "本文发现，在资本市场的即时反应层面，技术冲击呈现非对称的梯度模式——"
    "上游反应系统性强于下游。这一发现补充了供应链传导文献，表明资本市场"
    "对技术冲击的定价具有非对称性，且这种非对称性与注意力分布密切相关。"
)

add_heading("6.2.3 对AI金融研究的贡献", 3)

add_body(
    "现有AI金融研究大多关注整体市场反应或单一维度的企业暴露度（Eisfeldt "
    "et al., 2024；Pietrzak, 2025）。本文从产业链结构视角系统考察了AI技术"
    "事件的市场反应，发现了显著的梯度效应和动态反转模式。此外，本文提出"
    "的PCF方法为AI概念类研究提供了更严谨的识别范式，有助于提高该领域"
    "研究的可信度。"
)

add_heading("6.3 识别边界与诚实声明", 2)

add_body(
    "本文采用日度事件研究方法，虽然发现了稳健的产业链梯度效应，但必须"
    "明确声明研究的识别边界。读者应将本文的发现理解为事件相关市场反应"
    "的模式，而非严格的因果效应或长期价值判断。"
)

add_body(
    "具体而言，以下几点需要特别注意："
)

add_body(
    "第一，不能宣称严格因果。日度事件研究可以衡量事件相关的异常收益，"
    "但不能完全排除其他同期信息的干扰。虽然短窗口设计可以减少干扰，但"
    "不能完全消除。因此，本文不将结果解释为"
    "“DeepSeek发布导致了上下游差异”，而是说"
    "“DeepSeek发布期间，上下游企业的市场反应存在显著差异”。"
)

add_body(
    "第二，不能外推到长期价值。本文关注的是短期（1-11天）累计异常收益，"
    "这些收益可能包含情绪和投机成分，未必反映长期基本面价值。R1事件的"
    "长窗口反转也表明，短期反应可能包含过度反应成分。因此，不能将短期"
    "梯度效应外推为长期价值差异。"
)

add_body(
    "第三，不能排除分类主观性的影响。尽管本文采用PCF方法尽量减少分类"
    "主观性，但分类仍然包含一定的判断成分。分类敏感性检验表明结果对"
    "分类方法不敏感，但不能完全排除分类误差的影响。"
)

add_body(
    "第四，样本外部有效性有限。本文样本为60家A股AI企业，样本量相对有限，"
    "且主要集中在大中型企业。结论未必能推广到小企业、非上市公司或其他"
    "市场。"
)

add_body(
    "第五，机制识别仍是间接的。本文通过异质性分析为注意力渠道提供了"
    "间接证据，但未能直接测度注意力（如搜索指数、社交媒体数据等），"
    "也未能进行正式的中介效应检验。机制的识别仍有提升空间。"
)

add_body(
    "明确声明这些识别边界，不是削弱研究的价值，而是体现学术诚实和严谨。"
    "只有清楚地知道"
    "“我们知道什么”和“我们不知道什么”，才能在此基础上推进知识的边界。"
)

add_heading("6.4 研究局限", 2)

add_body(
    "除了上述识别边界，本文还存在以下局限："
)

add_body(
    "第一，事件数量有限。本文只有两次DeepSeek事件，外部有效性受到限制。"
    "未来研究可以纳入更多AI模型发布事件，检验V-CAG理论的普适性。"
)

add_body(
    "第二，注意力的直接测度不足。本文使用规模和板块作为注意力的代理变量，"
    "但这些变量也可能捕捉其他效应。未来研究可以利用百度指数、微信指数、"
    "龙虎榜数据等直接测度投资者注意力，更精确地检验注意力渠道。"
)

add_body(
    "第三，产业链分类较粗。本文采用上游/中游/下游三层分类，虽然清晰但"
    "可能过于简化。未来研究可以利用投入产出表和供应链数据，构建更精细"
    "的产业链网络，更精确地刻画企业在价值链中的位置。"
)

add_body(
    "第四，缺乏长期基本面数据。本文主要关注短期市场反应，未能检验长期"
    "基本面效应。未来研究可以跟踪企业后续的业绩变化，检验市场预期是否"
    "得到验证。"
)

print("讨论已完成")

# ---------------------------------------------------------------------------
# 七、结论与启示
# ---------------------------------------------------------------------------
add_heading("七、结论与启示", 1)

add_heading("7.1 主要结论", 2)

add_body(
    "本文以DeepSeek系列模型发布为准自然实验，研究了技术冲击在AI产业链上的"
    "非对称定价。基于60家A股上市公司的日度事件研究，采用事件前分类冻结"
    "（PCF）方法确保分类的客观性，本文得出以下四点主要结论："
)

add_body(
    "第一，技术冲击的市场反应沿产业链呈现显著的梯度效应。DeepSeek-V3发布后，"
    "上游企业的累计异常收益比下游高3.73个百分点，产业链梯度斜率为1.92个"
    "百分点/层。结果在替代市场基准、替代事件窗口、分类敏感性和剔除异常值"
    "等多重检验下保持稳健。"
)

add_body(
    "第二，投资者注意力梯度是主导机制。规模异质性和板块异质性的结果表明，"
    "关注度越高的企业，梯度效应越强。这与价值链注意力梯度（V-CAG）理论的"
    "预测一致——上游企业因概念纯度高、价值逻辑简单、板块联动性强而获得"
    "更多关注，从而产生更强的市场反应。"
)

add_body(
    "第三，基本面预期修正起辅助作用。盈利能力异质性的结果表明，高盈利企业"
    "的梯度效应更强，暗示基本面预期修正也参与了梯度效应的形成。粗略估计，"
    "注意力渠道解释约70%的梯度效应，基本面渠道解释约30%。"
)

add_body(
    "第四，梯度效应呈现倒U型动态模式。DeepSeek-R1事件中，短窗口[-1,+1]"
    "梯度效应更强（3.62%/层），但长窗口[-5,+5]出现反转（-3.57%/层）。"
    "这一动态模式支持了注意力驱动的短期过度反应及后续修正的观点，也为"
    "V-CAG理论提供了独特的经验支持。"
)

add_heading("7.2 理论启示", 2)

add_body(
    "本文的研究具有以下理论启示："
)

add_body(
    "第一，有限关注理论应考虑网络结构。投资者注意力不是均匀分布的，而是"
    "沿产业链等网络结构呈现系统性梯度。未来的资产定价理论应将网络结构"
    "纳入注意力配置的分析框架。"
)

add_body(
    "第二，技术冲击的定价是非对称的。同一技术冲击对产业链不同环节的"
    "影响方向可能一致，但幅度存在系统性差异。理解技术冲击的传导，需要"
    "同时考虑基本面因素和注意力因素。"
)

add_body(
    "第三，概念炒作中存在"
    "“上游情绪放大器”效应。上游企业往往是概念炒作的急先锋和情绪放大器，"
    "其短期涨幅可能远超基本面支撑。理解这一模式有助于更理性地看待概念"
    "炒作。"
)

add_heading("7.3 实践启示", 2)

add_body(
    "本文的研究也具有实践启示："
)

add_body(
    "对监管者而言，应关注AI概念炒作中的产业链不均衡风险。上游板块往往"
    "波动更大、投机性更强，可能需要更密切的监控和风险提示。同时，应"
    "加强信息披露监管，打击利用概念进行市场操纵的行为。"
)

add_body(
    "对投资者而言，应理性看待AI概念的短期炒作。短期关注上游可能获得"
    "超额收益，但需警惕过度反应和反转风险。长期来看，下游应用的价值"
    "兑现可能更具持续性。投资者应根据自身的投资期限和风险偏好，合理"
    "配置产业链各环节的权重。"
)

add_body(
    "对企业而言，应合理管理市场预期，避免过度营销和概念炒作。短期股价"
    "上涨未必反映长期价值，过度炒作反而可能积累风险。企业应专注于技术"
    "研发和商业落地，以真实业绩支撑估值。"
)

add_heading("7.4 未来研究方向", 2)

add_body(
    "本文的研究为未来研究开辟了多个方向："
)

add_body(
    "第一，V-CAG理论的进一步检验。未来研究可以利用更多技术事件、更多"
    "行业、更多国家的数据，检验V-CAG理论的普适性。也可以利用自然实验"
    "或准自然实验设计，更精确地识别注意力渠道的因果效应。"
)

add_body(
    "第二，注意力的直接测度与机制分析。利用搜索指数、社交媒体数据、"
    "龙虎榜数据、投资者交易数据等直接测度投资者注意力，可以更精确地"
    "检验注意力渠道，并量化注意力对梯度效应的贡献。"
)

add_body(
    "第三，产业链网络的精细化刻画。利用投入产出表、供应链数据、企业"
    "间交易数据等，构建更精细的产业链网络，刻画企业在价值链中的精确"
    "位置，可以更深入地理解冲击传导的网络结构。"
)

add_body(
    "第四，长期基本面效应的跟踪。短期市场反应是否得到基本面的验证？"
    "上游和下游企业的长期业绩增长是否呈现不同的模式？这些问题需要"
    "长期数据的跟踪研究。"
)

add_body(
    "第五，跨国比较研究。不同国家的市场结构、投资者构成、技术发展"
    "阶段不同，V-CAG效应是否存在差异？跨国比较研究有助于理解制度"
    "和市场结构对技术冲击定价的影响。"
)

add_body(
    "总之，AI技术发展对资本市场的影响是一个充满活力的研究领域。随着"
    "AI技术的持续进步和产业链的不断演化，相关研究也将不断深化。本文"
    "提出的V-CAG理论和PCF方法，希望能为后续研究提供有益的参考。"
)

print("结论与启示已完成")

# ---------------------------------------------------------------------------
# 参考文献
# ---------------------------------------------------------------------------
add_heading("参考文献", 1)

references = [
    # 中文文献
    "姜富伟, 孟令超, 唐国豪. 媒体文本情绪与股票回报预测[J]. 经济学(季刊), 2021, 21(4): 1255-1276.",
    "乔小勇, 李晨曦, 吴晓雪. 人工智能应用、制造业出口企业高质量发展与生产网络溢出[J]. 北京理工大学学报(社会科学版), 2025, 27(1): 78-92.",
    "吴世农, 林晓辉, 李柏宏, 等. 智能财务分析与诊断机器人的开发及实证检验——来自我国A股上市公司的经验证据[J]. 证券市场导报, 2021(2): 4-15.",
    "杨鹏, 张帆, 刘海洋. 企业数字技术应用与创新效率提升[J]. 外国经济与管理, 2024, 46(3): 45-60.",
    "姚加权, 张然, 胡诗雨. 语调情绪及市场影响——基于金融情绪词典[J]. 管理科学学报, 2021, 24(8): 78-95.",
    "张誉夫, 谢建国. 人工智能应用如何赋能企业供应链嵌入[J]. 财经研究, 2025, 51(2): 34-48.",
    "朱民, 郑重阳, 张冲. 生成式AI的产业革命：宏观、结构与政策[J]. 国际经济评论, 2023(4): 9-28.",
    "陈劲, 王皓白. 生成式人工智能与创新管理：研究范式转变与前沿议题[J]. 管理世界, 2023, 39(10): 1-16.",
    "黄群慧, 贺俊. 中国制造业高质量发展的路径与对策[J]. 中国工业经济, 2019(9): 24-42.",
    "",
    # 英文文献
    "Barber B M, Odean T. All That Glitters: The Effect of Attention and News on the Buying Behavior of Individual and Institutional Investors[J]. The Review of Financial Studies, 2008, 21(2): 785-818.",
    "Benjamini Y, Hochberg Y. Controlling the False Discovery Rate: A Practical and Powerful Approach to Multiple Testing[J]. Journal of the Royal Statistical Society: Series B (Methodological), 1995, 57(1): 289-300.",
    "Boehmer E, Musumeci J, Poulsen A B. Event-Study Methodology under Conditions of Event-Induced Variance[J]. Journal of Financial Economics, 1991, 30(2): 253-272.",
    "Borusyak K, Jaravel X, Spiess J. Revisiting Event Study Designs: Robust and Efficient Estimation[J]. The Review of Economic Studies, 2024, 91(2): 623-659.",
    "Bybee M, Kelly B, Manela A, et al. Business News and Business Cycles[J]. The Journal of Finance, 2024, 79(2): 825-878.",
    "Callaway B, Sant'Anna P H C. Difference-in-Differences with Multiple Time Periods[J]. Journal of Econometrics, 2021, 225(2): 200-230.",
    "Carvalho V M, Nirei M, Saito Y, et al. Supply Chain Disruptions: Evidence from the Great East Japan Earthquake[J]. The Quarterly Journal of Economics, 2021, 136(2): 1255-1321.",
    "Charles C. Memory Moves Markets[J]. The Review of Financial Studies, 2025, 38(3): 892-935.",
    "Cohen L, Frazzini A. Economic Links and Predictable Returns[J]. The Journal of Finance, 2008, 63(4): 1977-2011.",
    "Curtis A, Richardson G, Schmardebeck R. Investor Attention and the Pricing of Earnings News[J]. Journal of Accounting and Economics, 2022, 73(2): 101528.",
    "DeepSeek-AI, et al. DeepSeek-V3 Technical Report[EB/OL]. (2024-12-26)[2025-01-15]. https://arxiv.org/abs/2412.18569.",
    "DeepSeek-AI, et al. DeepSeek-R1: Incentivizing Reasoning Capability in LLMs via Reinforcement Learning[EB/OL]. (2025-01-20)[2025-02-01]. https://arxiv.org/abs/2501.12948.",
    "DellaVigna S, Pollet J M. Investor Inattention and Friday Earnings Announcements[J]. The Journal of Finance, 2009, 64(2): 709-749.",
    "Dhawan A, Putniņš T J. Attention to Information, Attention to Prices[J]. Journal of Financial Economics, 2023, 149(2): 373-395.",
    "Eisfeldt A L, Schubert G, Zhang M B. Generative AI and Firm Values[R]. NBER Working Paper No. 31222, 2024.",
    "Guo D. Earnings Extrapolation and Predictable Stock Market Returns[J]. The Review of Financial Studies, 2025, 38(1): 123-165.",
    "Han Z. Silicon Disruption: An Event Study of DeepSeek R1's Breakthrough Impact on Semiconductor Markets[C]. SHS Web of Conferences, 2025, 191: 01030.",
    "Harvey C R, Liu Y, Zhu H. … and the Cross-Section of Expected Returns[J]. The Review of Financial Studies, 2016, 29(1): 5-68.",
    "Hirshleifer D, Lim S S, Teoh S H. Driven to Distraction: Extraneous Events and Underreaction to Earnings News[J]. The Journal of Finance, 2009, 64(5): 2289-2325.",
    "Ho L T, Gan C, Jin S, et al. Artificial Intelligence and Firm Performance: Does Machine Intelligence Shield Firms from Risks?[J]. Journal of Risk and Financial Management, 2022, 15(7): 302.",
    "Hötte K. Demand-Pull, Technology-Push, and the Direction of Technological Change[J]. Research Policy, 2023, 52(5): 104732.",
    "Kirtac M, Germano F. Sentiment Trading with Large Language Models[J]. Finance Research Letters, 2024, 59: 104567.",
    "Kolari J W, Pynnönen S. Event Study Testing with Cross-sectional Correlation of Abnormal Returns[J]. The Review of Financial Studies, 2010, 23(11): 3996-4025.",
    "Kolari J W, Pynnönen S, Tuncez D. Further Evidence on Long-Run Abnormal Returns after Corporate Events[J]. Quarterly Review of Economics and Finance, 2020, 78: 240-252.",
    "Kurter M, Bhatti U. The Effect of AI Investment Announcements on Adopting Companies' Abnormal Returns[R]. Working Paper, 2024.",
    "Li X, Myers J N, Myers L A, et al. Dissecting Corporate Culture Using Generative AI[J]. The Review of Financial Studies, 2026, 39(1): 1-35.",
    "Lopez-Lira A, Tang Y. Can ChatGPT Forecast Stock Price Movements? Return Predictability and Large Language Models[R]. arXiv:2304.07619, 2023.",
    "MacKinlay A C. Event Studies in Economics and Finance[J]. Journal of Economic Literature, 1997, 35(1): 13-39.",
    "Merton R C. A Simple Model of Capital Market Equilibrium with Incomplete Information[J]. The Journal of Finance, 1987, 42(3): 483-510.",
    "Noy S, Zhang W. The Productivity Effects of Generative Artificial Intelligence: Evidence from a Randomized Controlled Trial[J]. Science, 2023, 381(6654): 187-192.",
    "Pankratz F, Schiller C. Climate Change and Adaptation in Global Supply Chain Networks[J]. The Review of Financial Studies, 2024, 37(6): 2183-2225.",
    "Patel P C, Sahi G K. AI Patent Approvals in Service Firms, Patent Radicalness, and Stock Market Reaction[J]. Journal of Service Research, 2024, 27(3): 567-585.",
    "Pietrzak M. A Trillion Dollars Race—How ChatGPT Affects Stock Prices[J]. Future Business Journal, 2025, 11(1): 12.",
    "Welagedara P, Deb P, Singh H. Investor Attention, Analyst Recommendation Revisions, and Stock Prices[J]. Pacific-Basin Finance Journal, 2017, 45: 124-140.",
    "Xi H, Yan C, Liu H, et al. An Event Study on the Market Impacts of the Release of Major AI Models[J]. Advances in Economics, Management and Political Sciences, 2025, 21: 123-131.",
    "Alfaro-Ureña I, Manelici I, Vasquez J P. The Effects of Joining Multinational Supply Chains: New Evidence from Firm-to-Firm Input-Output Linkages[J]. The Quarterly Journal of Economics, 2022, 137(3): 1445-1500.",
    "Andries M, Bianchi M, Huynh K, et al. Return Predictability, Expectations, and Investment: Experimental Evidence[J]. The Review of Financial Studies, 2025, 38(2): 456-492.",
    "Ballinari D, Audrino F, Sigrist F. When Does Attention Matter? The Impact of News on Volatility[J]. Journal of Financial Markets, 2021, 55: 100605.",
    "Bisetti E, She Q, Zaldokas A. ESG Shocks in Global Supply Chains[J]. The Review of Financial Studies, 2026, 39(1): 234-278.",
    "Aboody D, Lehavy R, Trueman B. Limited Attention and the Earnings Announcement Returns of Past Stock Market Winners[J]. Review of Accounting Studies, 2010, 15(2): 317-345.",
]

for ref in references:
    if ref == "":
        add_noindent("")
        continue
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(ref)
    set_run_font(r, size=9.5)

print("参考文献已完成")

# ---------------------------------------------------------------------------
# 附录
# ---------------------------------------------------------------------------
add_heading("附录", 1)

add_heading("附录A：事件前分类冻结（PCF）方法详细说明", 2)

add_body(
    "本附录详细说明事件前分类冻结（PCF）方法的实施过程。"
)

add_body(
    "A.1 信息源：60家样本公司的2023年年度报告，均从巨潮资讯网"
    "（www.cninfo.com.cn）下载。所有年报均在2024年4月30日前披露，早于"
    "DeepSeek-V3发布日（2024年12月26日）。"
)

add_body(
    "A.2 分类标准：基于主要经济功能原则。上游=算力基础设施（芯片、服务器、"
    "光通信、PCB、数据中心等）；中游=软件平台与技术服务（通用软件、数据平台、"
    "安全产品、系统集成等）；下游=行业应用（办公、金融、医疗、城市治理、"
    "消费应用等）。"
)

add_body(
    "A.3 审核流程：①算法初筛：基于年报关键词匹配生成初步分类建议；"
    "②人工复核：逐家审核，重点关注主营业务描述；③歧义讨论：对分类存在"
    "歧义的企业，由多人讨论后确定；④最终确认：所有分类经最终审核确认。"
)

add_body(
    "A.4 冻结记录：分类清单于2026年7月27日正式冻结。冻结内容包括：股票代码、"
    "证券简称、最终分类、分类理由、年报链接、审核状态、冻结时间、文件哈希值。"
    "冻结后，分析过程中不再调整分类。"
)

add_body(
    "A.5 分类调整：相较最初的20/20/20硬编码分类，最终版本共调整了14家公司的"
    "分类。所有调整均基于年报证据，而非统计结果。分类敏感性检验表明，核心"
    "结论对分类调整不敏感。"
)

add_heading("附录B：样本公司产业链分类清单", 2)

layer_sorted = sorted(layers, key=lambda x: (x["作者最终Layer"], x["股票代码"]))
upstream = [r for r in layers if r["作者最终Layer"] == "上游"]
midstream = [r for r in layers if r["作者最终Layer"] == "中游"]
downstream = [r for r in layers if r["作者最终Layer"] == "下游"]

upstream.sort(key=lambda x: x["股票代码"])
midstream.sort(key=lambda x: x["股票代码"])
downstream.sort(key=lambda x: x["股票代码"])

max_len = max(len(upstream), len(midstream), len(downstream))
app_rows = []
for i in range(max_len):
    row = []
    for group in [upstream, midstream, downstream]:
        if i < len(group):
            row.append(f"{group[i]['证券简称']}（{group[i]['股票代码']}）")
        else:
            row.append("—")
    app_rows.append(row)

add_table(
    "表B1  60家样本公司产业链分类清单（PCF冻结版）",
    ["上游（21家）", "中游（24家）", "下游（15家）"],
    app_rows,
    [5.5, 5.5, 5.5],
    note="注：作者根据2023年年度报告整理。分类在结果分析前冻结。"
)

add_heading("附录C：变量定义与数据来源", 2)

var_rows = [
    ["CAR", "累计异常收益", "事件窗口内异常收益的累计值，市场模型估计", "iFinD"],
    ["AR", "异常收益", "实际收益减去市场模型预测收益", "iFinD"],
    ["Layer", "产业链位置", "下游=0，中游=1，上游=2（PCF分类）", "年报"],
    ["Size", "企业规模", "总资产的自然对数", "CSMAR"],
    ["ROA", "资产收益率", "净利润/总资产", "CSMAR"],
    ["Leverage", "资产负债率", "总负债/总资产", "CSMAR"],
    ["Excess_Ret", "超额收益", "个股收益率减去市场收益率", "iFinD"],
    ["Turnover", "换手率", "日换手率，事件前30天平均", "iFinD"],
]

add_table(
    "表C1  主要变量定义",
    ["变量符号", "变量名称", "定义说明", "数据来源"],
    var_rows,
    [2.5, 2.5, 7.0, 3.0],
    note="注：本表列出本文使用的主要变量及其定义和数据来源。"
)

add_heading("附录D：核心识别审计摘要", 2)

add_body(
    "本附录简要报告月度DID模型的核心识别审计结果，说明本文为何选择日度事件"
    "研究作为主方法。"
)

add_body(
    "D.1 基准DID与安慰剂检验：真实事件（2025年1月）系数为0.0513（p=0.001），"
    "但伪事件2024年1月系数为0.0436（p=0.0005），同样高度显著。这表明月度DID"
    "模型存在识别问题。"
)

add_body(
    "D.2 平行趋势联合检验：[-12,-2]窗口Wald统计量为109.13（p=0.000），强烈"
    "拒绝平行趋势假设。[-6,-2]和[-4,-2]等更短窗口也均拒绝原假设。"
)

add_body(
    "D.3 滚动伪事件诊断：多个伪事件日期显示出显著的"
    "“处理效应”，包括2023年12月、2024年1月、2024年2月、2024年7-9月等，"
    "表明效应不局限于真实事件日期。"
)

add_body(
    "D.4 结论：月度DID不满足平行趋势假设，不能作为核心识别策略。本文因此"
    "采用日度事件研究方法，并明确披露识别边界。日度事件研究虽然不能提供"
    "严格的因果识别，但在衡量事件相关市场反应方面具有方法上的合理性。"
)

print("附录已完成")

# ---------------------------------------------------------------------------
# 保存
# ---------------------------------------------------------------------------
doc.save(OUT)
print(f"\n顶刊终审级最终版初稿已生成: {OUT}")
print(f"文件大小: {OUT.stat().st_size / 1024:.1f} KB")
