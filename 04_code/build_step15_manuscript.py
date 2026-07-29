# -*- coding: utf-8 -*-
"""步骤15：生成基于最终冻结Layer与日度事件研究的中文投稿版初稿。"""

from pathlib import Path
import sys
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"D:/thailand study/26_7_23paper")
OUT = (
    BASE / "05_output/发送/"
    "步骤15-投稿版初稿-DeepSeek事件与AI产业链非对称市场反应.docx"
)
MAIN_PATH = BASE / "05_output/revision_step14f/tables/table_j5_final_main_results.csv"
PAIR_PATH = BASE / "05_output/revision_step14f/tables/table_j3_final_pairwise_results.csv"
GRAD_PATH = BASE / "05_output/revision_step14f/tables/table_j4_final_gradient_results.csv"
LAYER_PATH = BASE / "05_output/revision_step14f/tables/table_j1_final_frozen_layer.csv"

main = pd.read_csv(MAIN_PATH, encoding="utf-8-sig")
pairs = pd.read_csv(PAIR_PATH, encoding="utf-8-sig")
gradients = pd.read_csv(GRAD_PATH, encoding="utf-8-sig")
layers = pd.read_csv(LAYER_PATH, encoding="utf-8-sig", dtype={"股票代码": str})

doc = Document()
section = doc.sections[0]
# narrative_proposal预设的“中国学术稿件”命名覆盖：A4、宋体、适度紧凑。
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.7)
section.right_margin = Cm(2.7)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)


def set_run_font(run, east_asia="宋体", latin="Times New Roman",
                 size=10.5, bold=False, color="000000"):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start),
                          ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(10.5)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.line_spacing = 1.35
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.first_line_indent = Cm(0.74)

for name, size, before, after in (
    ("Heading 1", 15, 16, 8),
    ("Heading 2", 13, 12, 6),
    ("Heading 3", 11.5, 8, 4),
):
    style = styles[name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(text)
    set_run_font(r, east_asia="黑体", latin="Arial", size=18, bold=True)


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_run_font(r, east_asia="宋体", size=10.5, color="555555")


def add_heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run(text)
    return p


def add_body(text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_noindent(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    set_run_font(r)
    r.font.italic = italic
    return p


def add_table(title, headers, rows, widths_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, east_asia="黑体", size=10.5, bold=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for index, (cell, header, width) in enumerate(
        zip(table.rows[0].cells, headers, widths_cm)
    ):
        cell.width = Cm(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "E7E6E6")
        set_cell_margins(cell)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run(str(header))
        set_run_font(run, east_asia="黑体", size=9, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row, widths_cm):
            cell.width = Cm(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=8.5)
    source = doc.add_paragraph()
    source.paragraph_format.first_line_indent = Cm(0)
    source.paragraph_format.space_before = Pt(3)
    source.paragraph_format.space_after = Pt(6)
    r = source.add_run("注：作者根据步骤14F最终冻结分类及事件研究结果整理。")
    set_run_font(r, size=8.5, color="555555")
    return table


add_title("DeepSeek发布事件与AI产业链的非对称市场反应")
add_subtitle("——基于A股上市公司的日度事件研究")
add_noindent("投稿版中文初稿（步骤15）")

add_heading("摘要", 1)
add_body(
    "生成式人工智能模型的突破如何沿产业链被资本市场定价，仍缺少基于企业经济功能和短窗口市场反应的系统证据。本文以DeepSeek-V3于2024年12月26日发布为核心事件，选取60家A股人工智能产业链上市公司，依据事件前公开的2023年年度报告，将企业按主要经济功能划分为上游、中游和下游，并采用市场模型估计日度异常收益与累计异常收益。以沪深300为主基准、[-1,+1]为主窗口，上游企业相对下游企业的累计异常收益高3.73个百分点（p=0.010；BH-FDR校正p=0.026）；将产业链位置编码为由下游至上游递增的有序变量后，每上移一个层级，累计异常收益提高1.92个百分点（p=0.004；BH-FDR校正p=0.006）。替换市场指数、扩大事件窗口以及使用DeepSeek-R1事件后，短窗口内的产业链梯度总体保持为正，但部分结果不能通过最保守的Bonferroni校正，且R1事件在较长窗口中出现衰减或反转。研究表明，技术模型发布引发的短期市场反应具有产业链位置差异；但这一证据属于事件相关市场反应，不足以单独识别长期价值创造或严格因果效应。"
)
add_noindent(
    "关键词：DeepSeek；AI产业链；事件研究；累计异常收益；产业链梯度；多重检验"
)

add_heading("一、引言", 1)
add_body(
    "大模型技术突破不仅改变人工智能产品的性能边界，也会重新组织资本市场对算力、软件平台和终端应用价值的判断。与一般企业新闻不同，基础模型发布同时影响产业链多个环节：芯片、服务器、印制电路板和光通信等上游企业提供计算与连接基础；软件平台、数据服务和系统集成企业承担技术扩散；办公、金融、医疗和城市治理等下游企业将模型嵌入具体场景。同一技术信息因而可能产生方向一致但幅度不同的价格反应。"
)
add_body(
    "现有注意力研究表明，投资者的信息获取范围有限，突发且显著的信息可能改变交易选择和短期价格压力（Merton, 1987；Barber and Odean, 2008）。但是，将这一逻辑用于AI产业链时存在两个经验困难。第一，企业是否属于上游或下游不能仅依据概念标签或事件后股价表现，而需要事件前可核查的主营业务证据。第二，单一技术事件并不天然构成标准双重差分设计；如果事件前趋势与时间安慰剂不能通过，显著的月度交互项也不能被直接解释为严格因果效应。"
)
add_body(
    "本文据此采用日度事件研究而非月度DID作为核心设计。研究以DeepSeek-V3发布日为主事件，先利用2024年12月26日前公开的2023年年度报告识别60家样本公司的主要经济功能，再通过市场模型计算不同事件窗口内的累计异常收益。分类过程在最终结果分析前冻结，并保留原始分类、调整理由、年报链接和文件哈希，以降低结果导向分类的风险。"
)
add_body(
    "本文提供三方面证据。第一，DeepSeek-V3发布后的短窗口市场反应沿产业链呈现正向梯度，上游企业反应强于下游企业。第二，该方向在上证综指、沪深300和深证成指三种基准下基本一致，并在较宽V3窗口中进一步增强。第三，DeepSeek-R1事件的短窗口也呈现正向梯度，但较长窗口发生衰减甚至反转，说明短期注意力反应不应被直接外推为长期基本面价值。"
)
add_body(
    "本文的贡献是提供一种可复核的产业链事件研究路径：用事件前年报确定企业经济功能，用预先锁定的事件、指数和窗口报告主结果，并同时披露BH-FDR与Bonferroni校正。相较于依赖概念标签和单一显著性水平的做法，这一路径更强调分类透明度、结果边界和可重复性。"
)

add_heading("二、理论分析与研究假设", 1)
add_heading("（一）技术发布、投资者注意力与短期异常收益", 2)
add_body(
    "有限关注意味着投资者不会同时处理全部公司信息。具有高传播度和低理解门槛的模型发布会集中吸引市场注意，并降低投资者搜索相关公司的成本。事件发生后，新增关注可能通过买入压力、预期修正和板块联动进入价格。不过，事件研究识别的是相对于市场模型预期的短期偏离，而不是对信息完全外生性的无条件证明。"
)
add_heading("（二）产业链位置与市场反应差异", 2)
add_body(
    "模型能力提升通常首先增加对训练、推理和部署基础设施的预期需求。芯片、服务器、光模块、PCB和数据中心等上游投入具有较清晰的产能与订单映射，市场更容易将模型突破转化为需求预期。中游平台和软件企业承担技术适配与扩散，其受益程度取决于产品兼容性和商业化能力。下游企业虽能获得应用叙事，但业务兑现往往依赖客户付费、场景改造和持续运营，因此短窗口反应可能较弱或分散。"
)
add_body(
    "基于上述逻辑，提出假设H1：DeepSeek模型发布后，上游企业的累计异常收益高于下游企业。提出假设H2：将产业链位置由下游、中游到上游依次编码为0、1和2时，累计异常收益随产业链位置上移而增加。"
)

add_heading("三、研究设计", 1)
add_heading("（一）事件选择与样本", 2)
add_body(
    "核心事件为DeepSeek-V3于2024年12月26日发布。补充事件为DeepSeek-R1于2025年1月20日发布。样本包括60家A股人工智能产业链上市公司。日度事件研究在每个事件、每个市场基准下均覆盖60家公司，共形成360个公司—事件—基准层面的CAR记录。"
)
add_heading("（二）事件前产业链分类", 2)
add_body(
    "产业链分类完全依据事件前公开信息。本文从巨潮资讯检索60家公司在2024年12月26日前披露的2023年年度报告正文，提取“主要业务”“主营业务”“主要产品”和“主要服务”等段落，再依据企业在AI价值链中的主要经济功能分类。最终冻结版本为layer_final_v1_2023_annual_report，上游21家、中游24家、下游15家。相较旧代码中的20/20/20硬编码名单，共调整14家公司。冻结表保存逐公司年报链接、分类理由、批准状态、冻结时间和SHA-256指纹。"
)

layer_counts = (
    layers.groupby("作者最终Layer")["股票代码"].count()
    .reindex(["上游", "中游", "下游"])
)
add_table(
    "表1  最终产业链分类分布",
    ["产业链位置", "公司数量", "经济功能"],
    [
        ["上游", int(layer_counts["上游"]), "芯片、服务器、光通信、PCB及算力基础设施"],
        ["中游", int(layer_counts["中游"]), "通用软件、数据平台、安全产品及系统集成"],
        ["下游", int(layer_counts["下游"]), "办公、金融、医疗、城市治理及消费应用"],
    ],
    [3.0, 3.0, 10.0],
)

add_heading("（三）异常收益计算", 2)
add_body(
    "本文使用标准市场模型估计正常收益（MacKinlay, 1997）。对公司i和交易日t，估计式为R_it=α_i+β_iR_mt+ε_it，其中R_mt分别使用沪深300、上证综指和深证成指收益。主规格使用沪深300。估计窗口内每家公司通常具有221个交易日观测。异常收益定义为实际收益减去市场模型预测收益，窗口累计异常收益为窗口内异常收益之和。本文报告[-1,+1]、[-3,+3]和[-5,+5]三个窗口，并预先将[-1,+1]设为主窗口。"
)
add_heading("（四）组间比较、梯度检验与多重校正", 2)
add_body(
    "上游与下游的均值差异采用Welch t检验，以允许两组方差不相等。梯度检验将下游、中游和上游分别编码为0、1和2，使用带HC1异方差稳健标准误的横截面回归。由于同时检验多个事件、指数、窗口和层级比较，本文报告Benjamini-Hochberg虚假发现率校正与Bonferroni校正（Benjamini and Hochberg, 1995）。前者用于控制发现集合中的预期错误比例，后者作为更保守的家族错误率边界。"
)

add_heading("四、实证结果", 1)
add_heading("（一）主规格结果", 2)
add_body(
    "表2报告DeepSeek-V3、沪深300基准和[-1,+1]窗口下的主结果。上游相对下游的CAR差异为0.0373，即3.73个百分点，原始p值为0.010，BH-FDR校正p值为0.026。产业链梯度斜率为0.0192，意味着从下游向上移动一个层级，CAR平均增加1.92个百分点；原始p值为0.004，BH-FDR校正p值为0.006。两项结果支持H1和H2。"
)
add_body(
    "不过，上游—下游差异的Bonferroni校正p值为0.548，梯度检验为0.076，均未达到5%水平。因此，本文将主结论限定为“在预设主规格和BH-FDR控制下存在显著产业链差异”，而不表述为“所有多重校正下均显著”。"
)
add_table(
    "表2  DeepSeek-V3事件的主规格结果",
    ["检验", "效应值", "原始p值", "BH-FDR", "Bonferroni"],
    [
        ["上游−下游", "0.0373", "0.0102", "0.0261", "0.5484"],
        ["产业链梯度", "0.0192", "0.0042", "0.0063", "0.0759"],
    ],
    [4.8, 2.8, 2.8, 2.8, 2.8],
)

add_heading("（二）替代市场基准与事件窗口", 2)
add_body(
    "使用上证综指和深证成指替代沪深300后，V3事件[-1,+1]窗口的上游—下游差异分别为0.0379和0.0373，方向与主规格一致。扩大至[-3,+3]后，沪深300基准下差异增至0.0895，梯度斜率为0.0455；扩大至[-5,+5]后，差异为0.1383，梯度斜率为0.0704。较宽窗口的结果更强，但同时更可能混入其他信息，因此只作为稳健性证据，不替代预设短窗口。"
)
add_heading("（三）DeepSeek-R1补充事件", 2)
add_body(
    "R1事件在[-1,+1]窗口呈现更强的短期梯度。以沪深300为基准，上游—下游差异为0.0727，梯度斜率为0.0362。两项结果均通过BH-FDR和Bonferroni校正。然而，当窗口扩大至[-5,+5]时，上游—下游差异和梯度转为负值。该反转说明R1的短窗口冲击可能包含集中注意力和即时价格压力，不能简单解释为持续性价值重估。"
)
add_table(
    "表3  主要稳健性与补充事件结果",
    ["事件", "基准", "窗口", "上游−下游", "梯度斜率", "解释"],
    [
        ["V3", "上证综指", "[-1,+1]", "0.0379", "0.0195", "方向一致"],
        ["V3", "深证成指", "[-1,+1]", "0.0373", "0.0192", "方向一致"],
        ["V3", "沪深300", "[-3,+3]", "0.0895", "0.0455", "较宽窗口增强"],
        ["V3", "沪深300", "[-5,+5]", "0.1383", "0.0704", "可能混入其他信息"],
        ["R1", "沪深300", "[-1,+1]", "0.0727", "0.0362", "短窗口显著"],
        ["R1", "沪深300", "[-5,+5]", "-0.0557", "-0.0357", "较长窗口反转"],
    ],
    [2.2, 2.8, 2.3, 2.5, 2.5, 4.0],
)

add_heading("（四）分类敏感性与月度DID诊断", 2)
add_body(
    "在原20/20/20硬编码分类下，V3主窗口的上游—下游差异为0.0415，产业链梯度为0.0208；使用事件前年报重新分类后，两者分别为0.0373和0.0192，方向与数量级接近。重新分类降低了部分显著性，但没有改变核心方向。这表明结果不是由个别明显错分公司完全驱动，同时也说明分类方法会影响推断强度。"
)
add_body(
    "作为识别诊断，本文曾使用月度面板和双向固定效应模型估计事件后交互项。修正收益口径后，基准交互项为0.0528（p<0.001），但2024年1月和2024年9月伪事件显著，事件前系数联合检验也拒绝零假设。因此，月度模型不满足将交互项解释为严格因果效应所需的关键识别条件。本文不以该模型支持主结论，而将其作为事件前趋势和同期行情影响的警示。"
)

add_heading("五、讨论", 1)
add_body(
    "产业链梯度可能反映市场将模型能力提升首先映射至可观察的算力和硬件需求。上游企业的产品边界相对明确，投资者较容易形成订单、资本开支和供需缺口预期；下游应用的商业化路径更依赖场景适配、客户付费和持续运营。因此，同一模型发布对不同经济功能企业产生不同幅度的短期价格反应。"
)
add_body(
    "但本文证据不能区分所有潜在机制。异常收益差异可能同时包含基本面预期、投资者注意力、板块轮动和风险偏好变化。R1结果在较长窗口中反转，进一步表明短期反应不等于长期现金流实现。若要识别机制，需要加入订单、盈利预测修正、搜索强度或投资者交易结构等独立数据。"
)
add_body(
    "多重检验结果也要求控制表述强度。BH-FDR适合在一组相关检验中控制预期错误发现比例，而Bonferroni在检验较多、结果相关时可能非常保守。本文同时报告两种校正，不以选择性标准替代完整披露。主规格在BH-FDR下显著、在Bonferroni 5%标准下不显著，应被理解为有支持但仍需外部事件复制的证据。"
)

add_heading("六、结论", 1)
add_body(
    "本文利用DeepSeek-V3发布事件，考察A股人工智能产业链公司的短期市场反应。基于60家公司事件前2023年年度报告建立的最终分类表，研究发现上游企业的[-1,+1]累计异常收益高于下游企业，且异常收益随产业链位置由下游向上游递增。替换市场指数和扩大V3窗口后方向总体一致，R1事件在短窗口提供补充支持，但较长窗口出现衰减或反转。"
)
add_body(
    "这些结果说明，资本市场对基础模型发布的即时反应并非均匀分布，而与企业在AI价值链中的主要经济功能相关。对投资者而言，概念相关性不能替代业务位置和商业化路径；对监管者而言，短期板块共振与持续价值重估需要区分；对企业而言，清晰披露产品在产业链中的功能有助于降低市场理解成本。"
)
add_body(
    "本文存在三项边界。第一，研究集中于中国A股60家公司和两次DeepSeek事件，外部有效性需要更多技术事件检验。第二，产业链分类虽依据事件前年报并已冻结，但跨环节经营仍可能造成测量误差。第三，事件研究能够描述短期异常收益，却不能单独排除全部同期信息，也不能证明长期价值创造。未来研究可结合高频交易、分析师预测修正和企业订单数据，进一步分解注意力与基本面渠道。"
)

add_heading("参考文献", 1)
refs = [
    "Barber, B. M., & Odean, T. (2008). All that glitters: The effect of attention and news on the buying behavior of individual and institutional investors. Review of Financial Studies, 21(2), 785–818.",
    "Benjamini, Y., & Hochberg, Y. (1995). Controlling the false discovery rate: A practical and powerful approach to multiple testing. Journal of the Royal Statistical Society: Series B, 57(1), 289–300.",
    "DeepSeek-AI et al. (2024). DeepSeek-V3 Technical Report. arXiv:2412.19437.",
    "DeepSeek-AI et al. (2025). DeepSeek-R1: Incentivizing reasoning capability in LLMs via reinforcement learning. arXiv:2501.12948.",
    "MacKinlay, A. C. (1997). Event studies in economics and finance. Journal of Economic Literature, 35(1), 13–39.",
    "Merton, R. C. (1987). A simple model of capital market equilibrium with incomplete information. Journal of Finance, 42(3), 483–510.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref)
    set_run_font(r, size=9.5)

# 页眉页脚
header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
header.paragraph_format.first_line_indent = Cm(0)
hr = header.add_run("DeepSeek发布事件与AI产业链的非对称市场反应")
set_run_font(hr, size=8.5, color="666666")
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.paragraph_format.first_line_indent = Cm(0)
fr = footer.add_run("步骤15投稿版初稿")
set_run_font(fr, size=8.5, color="777777")

# 防止表头跨页后丢失
for table in doc.tables:
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

doc.save(OUT)
print(OUT)

